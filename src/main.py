import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import sqlite3
from datetime import datetime
import openpyxl
import locale
import os
import glob


# --- Установка русского языка ---
locale.setlocale(locale.LC_TIME, 'russian')

# --- Получение пути до файла с бд ---
def get_db_path():
    current_dir = os.path.dirname(os.path.abspath(__file__))
    db_path = os.path.join(current_dir, "hotel.db")
    return db_path

# --- Класс, описывающий создание окна приложения ---
class App(tk.Tk):
    def __init__(self, title, size):    
        super().__init__()    
        self.title(title)    
        self.geometry(f'{size[0]}x{size[1]}') 
        self.minsize(size[0], size[1])
        
        # Создание фреймов
        self.menu = Menu(self)
        self.spravka = Spravka(self)
        self.about = About(self)
        self.clients = Clients(self)
        self.rent = Rent(self)
        self.room = Room(self)
        self.report = Report(self)
        self.payment = Payment(self)
        self.admin = AdminPanel(self)
        self.login = Login(self)
        
        # Передача ссылок в Menu для переключения фреймов
        self.menu.set_login_frame(self.login)
        self.menu.set_clients_frame(self.clients)
        self.menu.set_rent_frame(self.rent)
        self.menu.set_room_frame(self.room)
        self.menu.set_spravka_frame(self.spravka)
        self.menu.set_about_frame(self.about)
        self.menu.set_report_frame(self.report)
        self.menu.set_payment_frame(self.payment)
        self.menu.set_admin_frame(self.admin)
        
        # Настройка стиля
        style = ttk.Style(self)
        style.theme_use('clam')  # alt / default
        style.configure('TButton', foreground='white', background="#4D77A5", width=20)
        style.map('TButton', 
                background=[('active', "#35A7FF")],
                foreground=[('active', 'white')],
                font='Bold')
        # Сокрытие всех кнопок кроме "Вход" и "О программе" при старте
        self.menu.hide_main_buttons()
        self.mainloop()


# --- Класс, описывающий создание фрейма КЛИЕНТЫ ---
class Clients(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        testLbl = ttk.Label(self, background='#35A7FF')
        testLbl.pack(expand=True, fill='both')
        self.place(relx=0.3, y=0, relwidth=0.7, relheight=1)
        self.create_widgets()
        self.fetch_data()
        self.display_data()
        self.selected_item = None
        self.blacklist_data = []
        self.load_blacklist()
        
    # Создание виджетов
    def create_widgets(self):
        # Создание notebook для вкладок
        self.notebook = ttk.Notebook(self)
        self.notebook.place(relx=0.05, rely=0.05, relwidth=0.9, relheight=0.9)
        
        # Вкладка для всех клиентов
        self.clients_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.clients_tab, text="Все клиенты")
        self.create_clients_tab()
        
        # Вкладка для черного списка
        self.blacklist_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.blacklist_tab, text="Черный список")
        self.create_blacklist_tab()
        
        # Вкладка для истории бронирований
        self.history_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.history_tab, text="История")
        self.create_history_tab()

    # Создание вкладки История
    def create_history_tab(self):
        # Фрейм для поиска клиентов
        search_frame = ttk.Frame(self.history_tab)
        search_frame.pack(fill='x', padx=5, pady=5)
        
        # Поля для поиска клиентов
        ttk.Label(search_frame, text="Поиск клиента:").pack(side='left', padx=5)
        
        self.history_client_search_field = ttk.Combobox(search_frame, 
                                                     values=["ФИО", "Контакт", "Паспорт"], 
                                                     width=15, state='readonly')
        self.history_client_search_field.pack(side='left', padx=5)
        self.history_client_search_field.current(0)
        
        self.history_client_search_entry = ttk.Entry(search_frame)
        self.history_client_search_entry.pack(side='left', padx=5, expand=True, fill='x')
        
        ttk.Button(search_frame, text="Найти", command=self.search_history_clients).pack(side='left', padx=5)
        ttk.Button(search_frame, text="Сброс", command=self.reset_history_client_search).pack(side='left', padx=5)
        
        # Таблица клиентов для истории
        self.history_clients_tree = ttk.Treeview(
            self.history_tab, 
            columns=("ID", "ФИО", "Контакт", "Паспорт"), 
            show="headings",
            height=5
        )
        
        # Настройка колонок
        columns = {
            "ID": {"width": 50, "anchor": "center"},
            "ФИО": {"width": 200, "anchor": "w"},
            "Контакт": {"width": 120, "anchor": "center"},
            "Паспорт": {"width": 120, "anchor": "center"}
        }
        
        for col, params in columns.items():
            self.history_clients_tree.heading(col, text=col)
            self.history_clients_tree.column(col, **params)
        
        self.history_clients_tree.pack(fill='x', padx=5, pady=5)
        self.history_clients_tree.bind('<<TreeviewSelect>>', self.on_history_client_select)
        
        # Таблица истории бронирований выбранного клиента
        ttk.Label(self.history_tab, text="История бронирований:").pack(pady=(10, 0))
        
        self.history_bookings_tree = ttk.Treeview(
            self.history_tab, 
            columns=("ID", "Номер", "Заезд", "Выезд", "Сумма", "Статус"), 
            show="headings"
        )
        
        # Настройка колонок
        booking_columns = {
            "ID": {"width": 50, "anchor": "center"},
            "Номер": {"width": 80, "anchor": "center"},
            "Заезд": {"width": 100, "anchor": "center"},
            "Выезд": {"width": 100, "anchor": "center"},
            "Сумма": {"width": 100, "anchor": "e"},
            "Статус": {"width": 120, "anchor": "center"}
        }
        
        for col, params in booking_columns.items():
            self.history_bookings_tree.heading(col, text=col)
            self.history_bookings_tree.column(col, **params)
        
        self.history_bookings_tree.pack(expand=True, fill='both', padx=5, pady=5)
        
    # Поиск клиентов для просмотра истории
    def search_history_clients(self):    
        search_field = self.history_client_search_field.get()
        search_text = self.history_client_search_entry.get().strip()
        
        if not search_text:
            self.reset_history_client_search()
            return
            
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        
        if search_field == "ФИО":
            cursor.execute("SELECT id, name, contact, passport FROM Clients WHERE name LIKE ?", 
                          (f"%{search_text}%",))
        elif search_field == "Контакт":
            cursor.execute("SELECT id, name, contact, passport FROM Clients WHERE contact LIKE ?", 
                          (f"%{search_text}%",))
        elif search_field == "Паспорт":
            cursor.execute("SELECT id, name, contact, passport FROM Clients WHERE passport LIKE ?", 
                          (f"%{search_text}%",))
            
        clients = cursor.fetchall()
        conn.close()
        
        # Очищение таблицы
        for row in self.history_clients_tree.get_children():
            self.history_clients_tree.delete(row)
        
        # Заполнение результатами поиска
        for client in clients:
            self.history_clients_tree.insert("", "end", values=client)
            
    def reset_history_search(self):
        """Сброс поиска в истории"""
        self.history_search_entry.delete(0, tk.END)
        self.load_client_history()
    
    def reset_history_client_search(self):
        """Сброс поиска клиентов"""
        self.history_client_search_entry.delete(0, tk.END)
        
        # Загрузка всех клиентов
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute("SELECT id, name, contact, passport FROM Clients")
        clients = cursor.fetchall()
        conn.close()
        
        # Очищение таблицы
        for row in self.history_clients_tree.get_children():
            self.history_clients_tree.delete(row)
        
        # Заполние данными
        for client in clients:
            self.history_clients_tree.insert("", "end", values=client)
        
        # Очищение таблицы бронирований
        for row in self.history_bookings_tree.get_children():
            self.history_bookings_tree.delete(row)
            
    # Обработчик выбора клиента в истории
    def on_history_client_select(self, event):
        
        selected_item = self.history_clients_tree.selection()
        if not selected_item:
            return
            
        client_id = self.history_clients_tree.item(selected_item[0])['values'][0]
        self.load_client_bookings_history(client_id)

    # Загружает историю бронирований для выбранного клиента
    def load_client_bookings_history(self, client_id):
        
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT 
                r.id, 
                rm.room_number, 
                r.checkin_date, 
                r.checkout_date, 
                r.total_price,
                CASE 
                    WHEN r.checkout_date < date('now') THEN 'Завершено'
                    WHEN r.checkin_date > date('now') THEN 'Предстоящее'
                    ELSE 'Активное'
                END as status
            FROM Reservations r
            JOIN Rooms rm ON r.room_id = rm.id
            WHERE r.client_id = ?
            ORDER BY r.checkin_date DESC
        """, (client_id,))
        
        bookings = cursor.fetchall()
        conn.close()
        
        # Очищение таблицы
        for row in self.history_bookings_tree.get_children():
            self.history_bookings_tree.delete(row)
        
        # Заполнение данными
        for booking in bookings:
            res_id, room_number, checkin_date, checkout_date, total_price, status = booking
            self.history_bookings_tree.insert("", "end", values=(
                res_id,
                room_number,
                checkin_date,
                checkout_date,
                f"{total_price:.2f} руб.",
                status
            ))
            
    def create_clients_tab(self):
        # Фрейм для поиска
        search_frame = ttk.Frame(self.clients_tab)
        search_frame.pack(fill='x', padx=5, pady=5)
        
        # Поля для поиска
        ttk.Label(search_frame, text="Поиск:").pack(side='left', padx=5)
        
        self.search_field = ttk.Combobox(search_frame, values=["ФИО", "Контакт", "Паспорт", "Дата рождения"])
        self.search_field.pack(side='left', padx=5)
        self.search_field.current(0)
        
        self.search_entry = ttk.Entry(search_frame)
        self.search_entry.pack(side='left', padx=5, expand=True, fill='x')
        
        ttk.Button(search_frame, text="Найти", command=self.search_clients).pack(side='left', padx=5)
        ttk.Button(search_frame, text="Сброс", command=self.reset_search).pack(side='left', padx=5)
        
        # Кнопки управления
        button_frame = ttk.Frame(self.clients_tab)
        button_frame.pack(fill='x', padx=5, pady=5)
        
        # Создаем кнопок
        buttons = [
            ('Новый клиент', self.open_add_client_window),
            ('Изменить', self.open_edit_client_window),
            ('Удалить', self.delete_client),
            ('В черный список', self.add_to_blacklist),
            ('Обновить', self.refresh_data)
        ]
        
        for i, (text, command) in enumerate(buttons):
            btn = ttk.Button(button_frame, text=text, command=command)
            btn.grid(row=0, column=i, padx=2, sticky='ew')
            button_frame.columnconfigure(i, weight=1)  # изменение ширины колонок
        
        # Таблица клиентов
        self.columns = ("ID", "ФИО", "Контакт", "Паспорт", "Дата рождения")
        self.tree = ttk.Treeview(self.clients_tab, columns=self.columns, show="headings")
        
        for col in self.columns:
            self.tree.heading(col, text=col)
            self.tree.column(col, anchor='c')
            
        # Настройка ширины колонок с учетом возможного масштабирования
        self.tree.column("#1", width=50, stretch=False)
        self.tree.column("#2", width=200, stretch=True)
        self.tree.column("#3", width=150, stretch=True)
        self.tree.column("#4", width=150, stretch=True)
        self.tree.column("#5", width=150, stretch=True)
        
        self.tree.pack(expand=True, fill='both', padx=5, pady=5)
        self.tree.bind('<<TreeviewSelect>>', self.on_tree_select)

    def create_blacklist_tab(self):
        # Фрейм для поиска
        search_frame = ttk.Frame(self.blacklist_tab)
        search_frame.pack(fill='x', padx=5, pady=5)
        
        # Поля для поиска
        ttk.Label(search_frame, text="Поиск:").pack(side='left', padx=5)
        
        self.blacklist_search_field = ttk.Combobox(search_frame, values=["ФИО", "Контакт", "Паспорт", "Дата рождения", "Причина"])
        self.blacklist_search_field.pack(side='left', padx=5)
        self.blacklist_search_field.current(0)
        
        self.blacklist_search_entry = ttk.Entry(search_frame)
        self.blacklist_search_entry.pack(side='left', padx=5, expand=True, fill='x')
        
        ttk.Button(search_frame, text="Найти", command=self.search_blacklist).pack(side='left', padx=5)
        ttk.Button(search_frame, text="Сброс", command=self.reset_blacklist_search).pack(side='left', padx=5)
        
        # Кнопки управления черным списком
        button_frame = ttk.Frame(self.blacklist_tab)
        button_frame.pack(fill='x', padx=5, pady=5)
        
        ttk.Button(button_frame, text='Удалить из списка', command=self.remove_from_blacklist).pack(side='left', padx=5)
        ttk.Button(button_frame, text='Обновить', command=self.load_blacklist).pack(side='right', padx=5)
        
        # Таблица черного списка
        self.blacklist_tree = ttk.Treeview(
            self.blacklist_tab, 
            columns=("ID", "ФИО", "Контакт", "Паспорт", "Дата рождения", "Причина"), 
            show="headings"
        )
        
        columns = {
            "ID": {"width": 50, "anchor": "center"},
            "ФИО": {"width": 150, "anchor": "center"},
            "Контакт": {"width": 120, "anchor": "center"},
            "Паспорт": {"width": 120, "anchor": "center"},
            "Дата рождения": {"width": 120, "anchor": "center"},
            "Причина": {"width": 300, "anchor": "w", "stretch": True}
        }
        
        for col, params in columns.items():
            self.blacklist_tree.heading(col, text=col)
            self.blacklist_tree.column(col, **params)
        
        # перенос строк
        style = ttk.Style()
        style.configure("Treeview", rowheight=30)
        
        self.blacklist_tree.pack(expand=True, fill='both', padx=5, pady=5)
        self.blacklist_tree.bind('<<TreeviewSelect>>', self.on_blacklist_select)

    def on_tree_select(self, event):
        self.selected_item = self.tree.selection()[0] if self.tree.selection() else None

    def on_blacklist_select(self, event):
        self.selected_blacklist_item = self.blacklist_tree.selection()[0] if self.blacklist_tree.selection() else None

    def search_blacklist(self):
        search_field = self.blacklist_search_field.get()
        search_text = self.blacklist_search_entry.get()
        
        if not search_text:
            self.load_blacklist()
            return
            
        filtered_data = []
        
        for row in self.blacklist_data:
            if search_field == "ФИО" and search_text.lower() in row[1].lower():
                filtered_data.append(row)
            elif search_field == "Контакт" and search_text.lower() in row[2].lower():
                filtered_data.append(row)
            elif search_field == "Паспорт" and search_text.lower() in row[3].lower():
                filtered_data.append(row)
            elif search_field == "Дата рождения" and search_text.lower() in row[4].lower():
                filtered_data.append(row)
            elif search_field == "Причина" and search_text.lower() in row[5].lower():
                filtered_data.append(row)
        
        self.display_filtered_blacklist(filtered_data)

    def display_filtered_blacklist(self, data):
        for row in self.blacklist_tree.get_children():
            self.blacklist_tree.delete(row)
        for row in data:
            self.blacklist_tree.insert("", "end", values=row)

    def reset_blacklist_search(self):
        self.blacklist_search_entry.delete(0, tk.END)
        self.load_blacklist()

    def fetch_data(self):
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM Clients WHERE id NOT IN (SELECT client_id FROM Blacklist)")
        self.rows = cursor.fetchall()
        conn.close()

    def display_data(self):
        for row in self.tree.get_children():
            self.tree.delete(row)
        for row in self.rows:
            self.tree.insert("", "end", values=row)

    def load_blacklist(self):
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute("""
            SELECT c.id, c.name, c.contact, c.passport, c.birthdate, b.reason 
            FROM Clients c
            JOIN Blacklist b ON c.id = b.client_id
        """)
        self.blacklist_data = cursor.fetchall()
        conn.close()
        
        self.display_blacklist()

    def display_blacklist(self):
        for row in self.blacklist_tree.get_children():
            self.blacklist_tree.delete(row)
        for row in self.blacklist_data:
            self.blacklist_tree.insert("", "end", values=row)

    def refresh_data(self):
        self.fetch_data()
        self.display_data()
        self.load_blacklist()

    def search_clients(self):
        search_field = self.search_field.get()
        search_text = self.search_entry.get()
        
        if not search_text:
            self.refresh_data()
            return
            
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        
        if search_field == "ФИО":
            cursor.execute("SELECT * FROM Clients WHERE name LIKE ?", (f"%{search_text}%",))
        elif search_field == "Контакт":
            cursor.execute("SELECT * FROM Clients WHERE contact LIKE ?", (f"%{search_text}%",))
        elif search_field == "Паспорт":
            cursor.execute("SELECT * FROM Clients WHERE passport LIKE ?", (f"%{search_text}%",))
        elif search_field == "Дата рождения":
            cursor.execute("SELECT * FROM Clients WHERE birthdate LIKE ?", (f"%{search_text}%",))
            
        self.rows = cursor.fetchall()
        conn.close()
        self.display_data()

    def reset_search(self):
        self.search_entry.delete(0, tk.END)
        self.refresh_data()

    
    
    def open_add_client_window(self):
        self.add_window = tk.Toplevel(self)
        self.add_window.title("Добавить нового клиента")
        self.add_window.geometry("400x250")
        self.add_window.resizable(False, False)
        
        # Ограничения длины для полей ввода
        max_fio_length = 100
        max_contact_length = 50
        max_passport_length = 20
        max_birthdate_length = 10  # ДД.ММ.ГГГГ - 10 символов
        
        # Функции валидации
        def validate_length(text, max_len):
            return len(text) <= max_len
        
        vcmd_fio = (self.add_window.register(lambda text: validate_length(text, max_fio_length))), '%P'
        vcmd_contact = (self.add_window.register(lambda text: validate_length(text, max_contact_length))), '%P'
        vcmd_passport = (self.add_window.register(lambda text: validate_length(text, max_passport_length))), '%P'
        vcmd_birthdate = (self.add_window.register(lambda text: validate_length(text, max_birthdate_length))), '%P'
        
        style = ttk.Style()
        style.configure('Add.TLabel', padding=5, font=('Arial', 10))
        style.configure('Add.TEntry', padding=5, font=('Arial', 10))
        
        main_frame = ttk.Frame(self.add_window, padding=10)
        main_frame.pack(fill='both', expand=True)
        
        ttk.Label(main_frame, text="ФИО:", style='Add.TLabel').grid(row=0, column=0, sticky='e', pady=5)
        self.fio_entry = ttk.Entry(main_frame, style='Add.TEntry', validate="key", validatecommand=vcmd_fio)
        self.fio_entry.grid(row=0, column=1, sticky='ew', padx=5, pady=5)
        
        ttk.Label(main_frame, text="Контакт:", style='Add.TLabel').grid(row=1, column=0, sticky='e', pady=5)
        self.contact_entry = ttk.Entry(main_frame, style='Add.TEntry', validate="key", validatecommand=vcmd_contact)
        self.contact_entry.grid(row=1, column=1, sticky='ew', padx=5, pady=5)
        
        ttk.Label(main_frame, text="Паспорт:", style='Add.TLabel').grid(row=2, column=0, sticky='e', pady=5)
        self.passport_entry = ttk.Entry(main_frame, style='Add.TEntry', validate="key", validatecommand=vcmd_passport)
        self.passport_entry.grid(row=2, column=1, sticky='ew', padx=5, pady=5)
        
        ttk.Label(main_frame, text="Дата рождения:", style='Add.TLabel').grid(row=3, column=0, sticky='e', pady=5)
        self.birthdate_entry = ttk.Entry(main_frame, style='Add.TEntry', validate="key", validatecommand=vcmd_birthdate)
        self.birthdate_entry.grid(row=3, column=1, sticky='ew', padx=5, pady=5)

        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=4, column=0, columnspan=2, pady=10, sticky='e')
        
        ttk.Button(button_frame, text="Добавить", command=self.add_client).pack(side='right', padx=5)
        ttk.Button(button_frame, text="Отмена", command=self.add_window.destroy).pack(side='right', padx=5)
        
        main_frame.columnconfigure(1, weight=1)
        self.fio_entry.focus_set()


    def add_client(self):
        # Получение данных из полей ввода
        data = {
            "ФИО": self.fio_entry.get().strip(),
            "Контактные данные": self.contact_entry.get().strip(),
            "Паспортные данные": self.passport_entry.get().strip(),
            "Дата рождения": self.birthdate_entry.get().strip()
        }
        
        errors = []
        
        # Проверка заполненности полей
        for field, value in data.items():
            if not value:
                errors.append(f"• Поле '{field}' не заполнено")
        
        # Проверка максимальной длины
        if len(data["ФИО"]) > 100:
            errors.append("• ФИО слишком длинное (максимум 100 символов)")
        if len(data["Контактные данные"]) > 50:
            errors.append("• Контактные данные слишком длинные (максимум 50 символов)")
        if len(data["Паспортные данные"]) > 20:
            errors.append("• Паспортные данные слишком длинные (максимум 20 символов)")
        
        # Проверка формата даты
        if data["Дата рождения"]:
            try:
                birthdate = datetime.strptime(data["Дата рождения"], "%d.%m.%Y")
                # Проверка возраста
                age = (datetime.now() - birthdate).days / 365
                if age < 18:
                    if not messagebox.askyesno("Подтверждение", 
                                            "Клиент младше 18 лет. Вы уверены, что хотите добавить этого клиента?"):
                        return
            except ValueError:
                errors.append("• Некорректный формат даты рождения (требуется ДД.ММ.ГГГГ)")
        
        # обработка ошибок
        if errors:
            error_message = "Исправьте следующие ошибки:\n\n" + "\n".join(errors)
            error_message += "\n\nПример корректных данных:\n"
            error_message += "• ФИО: Иванов Иван Иванович\n"
            error_message += "• Контактные данные: 79001234567\n"
            error_message += "• Паспортные данные: 1234567890\n"
            error_message += "• Дата рождения: 15.05.1985"
            messagebox.showerror("Ошибка ввода данных", error_message)
            return
        
        # Проверка уникальности паспорта
        try:
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            cursor.execute("SELECT id FROM Clients WHERE passport = ?", (data["Паспортные данные"],))
            if cursor.fetchone():
                messagebox.showerror("Ошибка", "Клиент с такими паспортными данными уже существует")
                conn.close()
                return
                
            # Добавление клиента
            cursor.execute(
                "INSERT INTO Clients (name, contact, passport, birthdate) VALUES (?, ?, ?, ?)",
                (data["ФИО"], data["Контактные данные"], data["Паспортные данные"], data["Дата рождения"])
            )
            conn.commit()
            conn.close()
            
            messagebox.showinfo("Успех", "Клиент успешно добавлен")
            self.refresh_data()
            self.add_window.destroy()
            
        except sqlite3.Error as e:
            messagebox.showerror("Ошибка базы данных", f"Не удалось добавить клиента: {str(e)}")
            if conn:
                conn.close()

        
        
    def open_edit_client_window(self):
        if not self.selected_item:
            messagebox.showwarning("Предупреждение", "Выберите клиента для редактирования")
            return
        
        selected = self.tree.item(self.selected_item)
        client_data = selected['values']
        
        self.edit_window = tk.Toplevel(self)
        self.edit_window.title("Редактировать клиента")
        self.edit_window.geometry("400x250")
        self.edit_window.resizable(False, False)
        
        # Ограничения длины для полей ввода
        max_fio_length = 100
        max_contact_length = 50
        max_passport_length = 20
        max_birthdate_length = 10
        
        # Функции валидации
        def validate_length(text, max_len):
            return len(text) <= max_len
        
        vcmd_fio = (self.edit_window.register(lambda text: validate_length(text, max_fio_length))), '%P'
        vcmd_contact = (self.edit_window.register(lambda text: validate_length(text, max_contact_length))), '%P'
        vcmd_passport = (self.edit_window.register(lambda text: validate_length(text, max_passport_length))), '%P'
        vcmd_birthdate = (self.edit_window.register(lambda text: validate_length(text, max_birthdate_length))), '%P'
        
        style = ttk.Style()
        style.configure('Edit.TLabel', padding=5, font=('Arial', 10))
        style.configure('Edit.TEntry', padding=5, font=('Arial', 10))
        
        main_frame = ttk.Frame(self.edit_window, padding=10)
        main_frame.pack(fill='both', expand=True)
        
        ttk.Label(main_frame, text="ФИО:", style='Edit.TLabel').grid(row=0, column=0, sticky='e', pady=5)
        self.edit_fio_entry = ttk.Entry(main_frame, style='Edit.TEntry', validate="key", validatecommand=vcmd_fio)
        self.edit_fio_entry.grid(row=0, column=1, sticky='ew', padx=5, pady=5)
        self.edit_fio_entry.insert(0, client_data[1])
        
        ttk.Label(main_frame, text="Контакт:", style='Edit.TLabel').grid(row=1, column=0, sticky='e', pady=5)
        self.edit_contact_entry = ttk.Entry(main_frame, style='Edit.TEntry', validate="key", validatecommand=vcmd_contact)
        self.edit_contact_entry.grid(row=1, column=1, sticky='ew', padx=5, pady=5)
        self.edit_contact_entry.insert(0, client_data[2])
        
        ttk.Label(main_frame, text="Паспорт:", style='Edit.TLabel').grid(row=2, column=0, sticky='e', pady=5)
        self.edit_passport_entry = ttk.Entry(main_frame, style='Edit.TEntry', validate="key", validatecommand=vcmd_passport)
        self.edit_passport_entry.grid(row=2, column=1, sticky='ew', padx=5, pady=5)
        self.edit_passport_entry.insert(0, client_data[3])
        
        ttk.Label(main_frame, text="Дата рождения:", style='Edit.TLabel').grid(row=3, column=0, sticky='e', pady=5)
        self.edit_birthdate_entry = ttk.Entry(main_frame, style='Edit.TEntry', validate="key", validatecommand=vcmd_birthdate)
        self.edit_birthdate_entry.grid(row=3, column=1, sticky='ew', padx=5, pady=5)
        self.edit_birthdate_entry.insert(0, client_data[4])

        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=4, column=0, columnspan=2, pady=10, sticky='e')
        
        ttk.Button(button_frame, text="Сохранить", command=self.edit_client).pack(side='right', padx=5)
        ttk.Button(button_frame, text="Отмена", command=self.edit_window.destroy).pack(side='right', padx=5)
        
        main_frame.columnconfigure(1, weight=1)
        self.edit_fio_entry.focus_set()


    def edit_client(self):
        client_id = self.tree.item(self.selected_item)['values'][0]
        fio = self.edit_fio_entry.get()
        contact = self.edit_contact_entry.get()
        passport = self.edit_passport_entry.get()
        birthdate = self.edit_birthdate_entry.get()
        
        if not all([fio, contact, passport, birthdate]):
            messagebox.showerror("Ошибка", "Все поля должны быть заполнены")
            return
            
        try:
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            cursor.execute("UPDATE Clients SET name=?, contact=?, passport=?, birthdate=? WHERE id=?",
                          (fio, contact, passport, birthdate, client_id))
            conn.commit()
            conn.close()
            self.refresh_data()
            self.edit_window.destroy()
            messagebox.showinfo("Успех", "Данные клиента успешно обновлены")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось обновить данные: {str(e)}")
            
        try:
            datetime.strptime(birthdate, "%d.%m.%Y")
        except ValueError:
            messagebox.showerror("Ошибка", "Некорректный формат даты. Используйте ДД.ММ.ГГГГ")
            return

    def delete_client(self):
        if not self.selected_item:
            messagebox.showwarning("Предупреждение", "Выберите клиента для удаления")
            return
            
        client_id = self.tree.item(self.selected_item)['values'][0]
        client_name = self.tree.item(self.selected_item)['values'][1]
        
        if messagebox.askyesno("Подтверждение", f"Вы действительно хотите удалить клиента {client_name}?"):
            try:
                conn = sqlite3.connect(get_db_path())
                cursor = conn.cursor()
                cursor.execute("DELETE FROM Clients WHERE id=?", (client_id,))
                conn.commit()
                conn.close()
                self.refresh_data()
                messagebox.showinfo("Успех", "Клиент успешно удален")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось удалить клиента: {str(e)}")

    def add_to_blacklist(self):
        if not self.selected_item:
            messagebox.showwarning("Предупреждение", "Выберите клиента для добавления в черный список")
            return
            
        client_id = self.tree.item(self.selected_item)['values'][0]
        client_name = self.tree.item(self.selected_item)['values'][1]
        
        self.blacklist_window = tk.Toplevel(self)
        self.blacklist_window.title("Добавить в черный список")
        self.blacklist_window.geometry("400x200")
        
        tk.Label(self.blacklist_window, text=f"Клиент: {client_name}").pack(pady=5)
        tk.Label(self.blacklist_window, text="Причина:").pack()
        
        # Функция валидации длины
        def validate_length(text):
            return len(text) <= 50
        
        vcmd = (self.blacklist_window.register(validate_length)), '%P'
        
        self.reason_entry = tk.Entry(self.blacklist_window, width=60, validate="key", validatecommand=vcmd)
        self.reason_entry.pack(pady=5)
        
        tk.Button(self.blacklist_window, text="Добавить", 
                command=lambda: self.confirm_add_to_blacklist(client_id)).pack(pady=5)

    def confirm_add_to_blacklist(self, client_id):
        reason = self.reason_entry.get().strip()
        
        if not reason:
            messagebox.showerror("Ошибка", "Укажите причину добавления в черный список")
            return
            
        try:
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            
            cursor.execute("SELECT 1 FROM Blacklist WHERE client_id=?", (client_id,))
            if cursor.fetchone():
                messagebox.showerror("Ошибка", "Этот клиент уже в черном списке")
                conn.close()
                return
                
            cursor.execute("INSERT INTO Blacklist (client_id, reason) VALUES (?, ?)", (client_id, reason))
            conn.commit()
            conn.close()
            
            self.blacklist_window.destroy()
            self.refresh_data()
            messagebox.showinfo("Успех", "Клиент добавлен в черный список")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось добавить клиента в черный список: {str(e)}")
    
    def remove_from_blacklist(self):
        if not self.selected_blacklist_item:
            messagebox.showwarning("Предупреждение", "Выберите клиента для удаления из черного списка")
            return
            
        client_id = self.blacklist_tree.item(self.selected_blacklist_item)['values'][0]
        client_name = self.blacklist_tree.item(self.selected_blacklist_item)['values'][1]
        
        if messagebox.askyesno("Подтверждение", 
                             f"Вы действительно хотите удалить клиента {client_name} из черного списка?"):
            try:
                conn = sqlite3.connect(get_db_path())
                cursor = conn.cursor()
                cursor.execute("DELETE FROM Blacklist WHERE client_id=?", (client_id,))
                conn.commit()
                conn.close()
                
                self.load_blacklist()
                messagebox.showinfo("Успех", "Клиент удален из черного списка")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось удалить клиента из черного списка: {str(e)}")


# --- Класс, описывающий создание фрейма БРОНИРОВАНИЕ ---
class Rent(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        self.style = ttk.Style()
        self.style.configure('Rent.TFrame', background='#f5f5f5')
        self.configure(style='Rent.TFrame')
        self.place(relx=0.3, y=0, relwidth=0.7, relheight=1)
        self.create_widgets()
        self.selected_item = None
        self.selected_clients = []
        self.fetch_data()

    def create_widgets(self):
        style = ttk.Style()
        style.configure('Rent.TFrame', background='#f0f0f0')
        style.configure('Rent.TLabel', background='#f0f0f0', font=('Arial', 10))
        style.configure('Rent.TButton', font=('Arial', 10))
        
        # Основной контейнер
        main_frame = ttk.Frame(self, style='Rent.TFrame')
        main_frame.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Фрейм для поиска
        search_frame = ttk.Frame(main_frame, style='Rent.TFrame')
        search_frame.pack(fill='x', pady=(0, 10))
        
        ttk.Label(search_frame, text="Поиск номеров:", style='Rent.TLabel').pack(side='left', padx=5)
        
        self.search_field = ttk.Combobox(search_frame, values=["Номер", "Класс", "Корпус", "Статус"], width=10)
        self.search_field.pack(side='left', padx=5)
        self.search_field.current(0)
        
        self.search_entry = ttk.Entry(search_frame, width=30)
        self.search_entry.pack(side='left', padx=5, expand=True, fill='x')
        
        ttk.Button(search_frame, text="Найти", command=self.search_rooms, style='Rent.TButton').pack(side='left', padx=5)
        ttk.Button(search_frame, text="Сброс", command=self.reset_search, style='Rent.TButton').pack(side='left', padx=5)
        
        # Блок выбора клиента с поиском
        client_frame = ttk.LabelFrame(main_frame, text=" Клиенты ", style='Rent.TFrame')
        client_frame.pack(fill='x', padx=5, pady=5)
        
        # Поиск клиентов
        client_search_frame = ttk.Frame(client_frame, style='Rent.TFrame')
        client_search_frame.pack(fill='x', pady=(0, 5))
        
        ttk.Label(client_search_frame, text="Поиск клиента:", style='Rent.TLabel').pack(side='left', padx=5)
        self.client_search_entry = ttk.Entry(client_search_frame)
        self.client_search_entry.pack(side='left', padx=5, expand=True, fill='x')
        self.client_search_entry.bind('<KeyRelease>', self.filter_clients)
        
        ttk.Label(client_frame, text="Выбрать клиента:", style='Rent.TLabel').pack(side='left', padx=5)
        self.client_combobox = ttk.Combobox(client_frame, width=25)
        self.client_combobox.pack(side='left', padx=5, pady=2)
        
        ttk.Button(client_frame, text="Добавить", command=self.add_client_to_booking,
                 style='Rent.TButton').pack(side='left', padx=5)
        
        ttk.Label(client_frame, text="Выбранные:", style='Rent.TLabel').pack(side='left', padx=5)
        self.selected_clients_listbox = tk.Listbox(client_frame, width=35, height=3)
        self.selected_clients_listbox.pack(side='left', padx=5, pady=2, fill='x', expand=True)
        
        # Фрейм для дат
        date_frame = ttk.Frame(main_frame, style='Rent.TFrame')
        date_frame.pack(fill='x', pady=5)
        
        ttk.Label(date_frame, text="Дата заезда:", style='Rent.TLabel').pack(side='left', padx=5)
        self.checkin_date = ttk.Entry(date_frame, width=12)
        self.checkin_date.pack(side='left', padx=5, pady=2)
        
        ttk.Label(date_frame, text="Дата выезда:", style='Rent.TLabel').pack(side='left', padx=5)
        self.checkout_date = ttk.Entry(date_frame, width=12)
        self.checkout_date.pack(side='left', padx=5, pady=2)
        
        # Фильтры для номеров
        filter_frame = ttk.Frame(main_frame, style='Rent.TFrame')
        filter_frame.pack(fill='x', pady=5)
        
        ttk.Label(filter_frame, text="Класс номера:", style='Rent.TLabel').pack(side='left', padx=5)
        self.class_combobox = ttk.Combobox(filter_frame, width=15)
        self.class_combobox.pack(side='left', padx=5)
        self.class_combobox.bind("<<ComboboxSelected>>", self.update_rooms_table)
        
        ttk.Label(filter_frame, text="Количество мест:", style='Rent.TLabel').pack(side='left', padx=5)
        self.places_combobox = ttk.Combobox(filter_frame, values=["1", "2", "3", "4"], width=5)
        self.places_combobox.pack(side='left', padx=5)
        self.places_combobox.bind("<<ComboboxSelected>>", self.update_rooms_table)
        
        # Таблица с номерами
        tree_frame = ttk.Frame(main_frame, style='Rent.TFrame')
        tree_frame.pack(fill='both', expand=True, pady=5)
        
        self.rooms_tree = ttk.Treeview(tree_frame, columns=("ID", "Номер", "Мест", "Класс", "Цена", "Корпус", "Статус"), 
                                     show="headings", height=8)
        
        for col in ("ID", "Номер", "Мест", "Класс", "Цена", "Корпус", "Статус"):
            self.rooms_tree.heading(col, text=col)
            self.rooms_tree.column(col, anchor='center', width=100)
        
        self.rooms_tree.pack(side='left', fill='both', expand=True)
        self.rooms_tree.bind('<<TreeviewSelect>>', self.on_room_select)
        
        scrollbar = ttk.Scrollbar(tree_frame, orient='vertical', command=self.rooms_tree.yview)
        scrollbar.pack(side='right', fill='y')
        self.rooms_tree.configure(yscrollcommand=scrollbar.set)
        
        # Кнопки управления
        button_frame = ttk.Frame(main_frame, style='Rent.TFrame')
        button_frame.pack(fill='x', pady=10)
        
        self.reserve_button = ttk.Button(button_frame, text="Забронировать", 
                                       command=self.create_reservation, state='disabled', style='Rent.TButton')
        self.reserve_button.pack(side='left', padx=5)
        
        ttk.Button(button_frame, text='Обновить', command=self.refresh_data, style='Rent.TButton').pack(side='left', padx=5)
        ttk.Button(button_frame, text="Удалить бронь", command=self.delete_reservation, style='Rent.TButton').pack(side='right', padx=5)
        
        # Таблица текущих броней
        reserv_tree_frame = ttk.Frame(main_frame, style='Rent.TFrame')
        reserv_tree_frame.pack(fill='both', expand=True, pady=(10, 0))
        
        self.reservations_tree = ttk.Treeview(reserv_tree_frame, 
                                            columns=("ID", "Номер", "Клиенты", "Заезд", "Выезд", "Сумма"), 
                                            show="headings", height=5)
        
        for col in ("ID", "Номер", "Клиенты", "Заезд", "Выезд", "Сумма"):
            self.reservations_tree.heading(col, text=col)
            self.reservations_tree.column(col, anchor='center', width=100)
        
        self.reservations_tree.pack(side='left', fill='both', expand=True)
        
        reserv_scrollbar = ttk.Scrollbar(reserv_tree_frame, orient='vertical', command=self.reservations_tree.yview)
        reserv_scrollbar.pack(side='right', fill='y')
        self.reservations_tree.configure(yscrollcommand=reserv_scrollbar.set)
    
        # Панель поиска бронирований
        reserv_search_frame = ttk.Frame(main_frame, style='Rent.TFrame')
        reserv_search_frame.pack(fill='x', pady=(10, 5))
        
        ttk.Label(reserv_search_frame, text="Поиск брони:", style='Rent.TLabel').pack(side='left', padx=5)
        
        self.reserv_search_field = ttk.Combobox(reserv_search_frame, 
                                             values=["Номер", "Клиент", "Дата заезда"], 
                                             width=12, state='readonly')
        self.reserv_search_field.current(0)
        self.reserv_search_field.pack(side='left', padx=5)
        
        self.reserv_search_entry = ttk.Entry(reserv_search_frame, width=30)
        self.reserv_search_entry.pack(side='left', padx=5, expand=True, fill='x')
        
        ttk.Button(reserv_search_frame, text="Найти", command=self.search_reservations, 
                 style='Rent.TButton').pack(side='left', padx=5)
        ttk.Button(reserv_search_frame, text="Сброс", command=self.reset_reserv_search, 
                 style='Rent.TButton').pack(side='left', padx=5)
    
    # Фильтрация списка клиентов по введенному тексту
    def filter_clients(self, event=None):
        
        search_text = self.client_search_entry.get().lower()
        filtered_clients = [c for c in self.clients if search_text in c[1].lower()]
        self.client_combobox['values'] = [f"{c[1]} (ID: {c[0]})" for c in filtered_clients]

    # Поиск бронирований по выбранному критерию
    def search_reservations(self):
        search_field = self.reserv_search_field.get()
        search_text = self.reserv_search_entry.get().strip()
        
        if not search_text:
            self.fetch_reservations()
            self.display_reservations()
            return
            
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        
        query = """
            SELECT r.id, rm.room_number, c.name, r.checkin_date, r.checkout_date, r.total_price
            FROM Reservations r
            JOIN Rooms rm ON r.room_id = rm.id
            JOIN Clients c ON r.client_id = c.id
            WHERE 1=1
        """
        
        if search_field == "Номер":
            query += " AND rm.room_number LIKE ?"
        elif search_field == "Клиент":
            query += " AND c.name LIKE ?"
        elif search_field == "Дата заезда":
            query += " AND r.checkin_date LIKE ?"
            
        cursor.execute(query, (f"%{search_text}%",))
        self.reservations = cursor.fetchall()
        conn.close()
        
        self.display_reservations()
        
    # Сброс поиска бронирований
    def reset_reserv_search(self):
        self.reserv_search_entry.delete(0, tk.END)
        self.fetch_reservations()
        self.display_reservations()

    # Загрузка клиентов с возможностью поиска
    def fetch_clients(self):
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute("SELECT id, name FROM Clients ORDER BY name")
        self.clients = cursor.fetchall()
        self.client_combobox['values'] = [f"{c[1]} (ID: {c[0]})" for c in self.clients]
        conn.close()
        
    # Загрузка всех необходимых данных
    def fetch_data(self):
        
        self.fetch_clients()
        self.fetch_rooms()
        self.update_rooms_table()
        self.fetch_reservations()
        self.display_reservations()

    # Поиск номеров по выбранному критерию
    def search_rooms(self):
        search_field = self.search_field.get()
        search_text = self.search_entry.get().strip()
        
        if not search_text:
            self.update_rooms_table()
            return
            
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        
        query = """
            SELECT r.id, r.room_number, r.places, rc.class_name, r.price, b.building_name, r.status
            FROM Rooms r
            JOIN RoomClasses rc ON r.class_id = rc.id
            JOIN Buildings b ON r.building_id = b.id
            WHERE 1=1
        """
        
        if search_field == "Номер":
            query += " AND r.room_number LIKE ?"
        elif search_field == "Класс":
            query += " AND rc.class_name LIKE ?"
        elif search_field == "Корпус":
            query += " AND b.building_name LIKE ?"
        elif search_field == "Статус":
            query += " AND r.status LIKE ?"
            
        cursor.execute(query, (f"%{search_text}%",))
        self.rooms = cursor.fetchall()
        conn.close()
        
        self.display_filtered_rooms()
        
    # Отображение отфильтрованных номеров
    def display_filtered_rooms(self):
        for row in self.rooms_tree.get_children():
            self.rooms_tree.delete(row)
        
        for room in self.rooms:
            self.rooms_tree.insert("", "end", values=room)
            
    # Сброс поиска
    def reset_search(self):
        self.search_entry.delete(0, tk.END)
        self.update_rooms_table()
    
    def fetch_clients(self):
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute("SELECT id, name FROM Clients")
        self.clients = cursor.fetchall()
        self.client_combobox['values'] = [f"{c[1]} (ID: {c[0]})" for c in self.clients]
        conn.close()

    def fetch_rooms(self):
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute("SELECT DISTINCT class_name FROM RoomClasses")
        self.classes = [row[0] for row in cursor.fetchall()]
        self.class_combobox['values'] = self.classes
        conn.close()

    def update_rooms_table(self, event=None):
        class_name = self.class_combobox.get()
        places = self.places_combobox.get()
        
        query = """
            SELECT r.id, r.room_number, r.places, rc.class_name, r.price, b.building_name, r.status
            FROM Rooms r
            JOIN RoomClasses rc ON r.class_id = rc.id
            JOIN Buildings b ON r.building_id = b.id
            WHERE 1=1
        """
        params = []
        
        if class_name:
            query += " AND rc.class_name = ?"
            params.append(class_name)
        
        if places:
            query += " AND r.places = ?"
            params.append(places)
        
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute(query, params)
        self.rooms = cursor.fetchall()
        
        # обновление таблицы
        for row in self.rooms_tree.get_children():
            self.rooms_tree.delete(row)
        
        for room in self.rooms:
            room_id, room_number, places, class_name, price, building_name, status = room
            
            # Проверка статуса номера
            if status in ["Требуется ремонт", "Требуется клининг"]:
                status_text = f"Недоступен ({status})"
            else:
                # Проверка бронирования только если статус не "недоступен"
                if self.checkin_date.get() and self.checkout_date.get():
                    try:
                        checkin = datetime.strptime(self.checkin_date.get(), "%d.%m.%Y").date()
                        checkout = datetime.strptime(self.checkout_date.get(), "%d.%m.%Y").date()
                        
                        cursor.execute("""
                            SELECT 1 FROM Reservations 
                            WHERE room_id = ? 
                            AND (
                                (checkin_date <= ? AND checkout_date > ?) 
                                OR (checkin_date < ? AND checkout_date >= ?)
                                OR (checkin_date >= ? AND checkout_date <= ?)
                            )
                        """, (room_id, checkin, checkin, checkout, checkout, checkin, checkout))
                        
                        if cursor.fetchone():
                            status_text = "Занят"
                        else:
                            status_text = "Свободен"
                    except ValueError:
                        status_text = "Свободен"
                else:
                    status_text = status
            
            self.rooms_tree.insert("", "end", values=(room_id, room_number, places, class_name, 
                                                price, building_name, status_text))
        
        conn.close()

    def add_client_to_booking(self):
        client_str = self.client_combobox.get()
        if not client_str:
            messagebox.showwarning("Предупреждение", "Выберите клиента")
            return
            
        if client_str in self.selected_clients:
            messagebox.showwarning("Предупреждение", "Этот клиент уже добавлен")
            return
            
        self.selected_clients.append(client_str)
        self.update_selected_clients_list()
        
        self.check_reservation_button_state()

    def remove_client_from_booking(self):
        if not self.selected_clients_listbox.curselection():
            messagebox.showwarning("Предупреждение", "Выберите клиента для удаления")
            return
            
        index = self.selected_clients_listbox.curselection()[0]
        self.selected_clients.pop(index)
        self.update_selected_clients_list()
        
        self.check_reservation_button_state()

    def update_selected_clients_list(self):
        self.selected_clients_listbox.delete(0, tk.END)
        for client in self.selected_clients:
            self.selected_clients_listbox.insert(tk.END, client)

    def check_reservation_button_state(self):
        if not self.selected_item:
            self.reserve_button['state'] = 'disabled'
            return
            
        room_data = self.rooms_tree.item(self.selected_item)['values']
        room_places = room_data[2]  # Количество мест в номере
        
        if len(self.selected_clients) > room_places:
            messagebox.showwarning("Предупреждение", 
                                 f"В номере только {room_places} мест, вы выбрали {len(self.selected_clients)} клиентов")
            self.reserve_button['state'] = 'disabled'
        elif len(self.selected_clients) > 0 and self.checkin_date.get() and self.checkout_date.get():
            self.reserve_button['state'] = 'normal'
        else:
            self.reserve_button['state'] = 'disabled'

    def on_room_select(self, event):
        self.selected_item = self.rooms_tree.selection()[0] if self.rooms_tree.selection() else None
        
        # Проверка количества мест в выбранном номере
        if self.selected_item:
            room_data = self.rooms_tree.item(self.selected_item)['values']
            room_places = room_data[2]
            
            if room_places == 1 and len(self.selected_clients) > 0:
                self.selected_clients = []
                self.update_selected_clients_list()
        
        self.check_reservation_button_state()

    def create_reservation(self):
        if not self.selected_item:
            messagebox.showwarning("Предупреждение", "Выберите номер для бронирования")
            return
            
        # Получение данные о номере
        room_id = self.rooms_tree.item(self.selected_item)['values'][0]
        room_number = self.rooms_tree.item(self.selected_item)['values'][1]
        room_places = self.rooms_tree.item(self.selected_item)['values'][2]
        room_status = self.rooms_tree.item(self.selected_item)['values'][6]
        
        # Проверка статуса номера
        if "Недоступен" in room_status or room_status in ["Требуется ремонт", "Требуется клининг"]:
            messagebox.showerror("Ошибка", f"Номер {room_number} недоступен для бронирования. Статус: {room_status}")
            return
        
        # Проверка количества клиентов
        if len(self.selected_clients) > room_places:
            messagebox.showwarning("Предупреждение", 
                                f"В номере только {room_places} мест, вы выбрали {len(self.selected_clients)} клиентов")
            return
            
        # Проверка даты
        if not self.checkin_date.get() or not self.checkout_date.get():
            messagebox.showerror("Ошибка", "Укажите даты заезда и выезда")
            return
        
        try:
            # Парсинг даты
            checkin = datetime.strptime(self.checkin_date.get(), "%d.%m.%Y").date()
            checkout = datetime.strptime(self.checkout_date.get(), "%d.%m.%Y").date()
            
            # корректность дат
            if checkout <= checkin:
                messagebox.showerror("Ошибка", "Дата выезда должна быть позже даты заезда")
                return
                
            # проверка, что дата заезда не в прошлом
            if checkin < datetime.now().date():
                messagebox.showerror("Ошибка", "Дата заезда не может быть в прошлом")
                return
                
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            
            # Проверка доступности номера на выбранные даты
            cursor.execute("""
                SELECT 1 FROM Reservations 
                WHERE room_id = ? 
                AND (
                    (checkin_date <= ? AND checkout_date > ?) 
                    OR (checkin_date < ? AND checkout_date >= ?)
                    OR (checkin_date >= ? AND checkout_date <= ?)
                )
            """, (room_id, checkin, checkin, checkout, checkout, checkin, checkout))
            
            if cursor.fetchone():
                messagebox.showerror("Ошибка", f"Номер {room_number} уже забронирован на выбранные даты")
                conn.close()
                return
                
            # Получение цены номера
            cursor.execute("SELECT price FROM Rooms WHERE id=?", (room_id,))
            price = cursor.fetchone()[0]
            
            # Рассчет суммы бронирования
            days = (checkout - checkin).days
            total = days * price
            
            # Создание брони для каждого клиента
            success_count = 0
            for client_str in self.selected_clients:
                try:
                    # получение ID клиента из строки вида "Имя (ID: 123)"
                    client_id = int(client_str.split("(ID: ")[1][:-1])
                    
                    # проверка клиента, что он не находится в черном списке
                    cursor.execute("SELECT 1 FROM Blacklist WHERE client_id=?", (client_id,))
                    if cursor.fetchone():
                        messagebox.showwarning("Предупреждение", 
                                            f"Клиент {client_str} находится в черном списке и не может быть заселен")
                        continue
                    
                    # Добавление бронирования
                    cursor.execute("""
                        INSERT INTO Reservations (room_id, client_id, checkin_date, checkout_date, total_price)
                        VALUES (?, ?, ?, ?, ?)
                    """, (room_id, client_id, checkin, checkout, total))
                    success_count += 1
                    
                except Exception as e:
                    messagebox.showerror("Ошибка", f"Не удалось создать бронь для клиента {client_str}: {str(e)}")
                    continue
            
            if success_count > 0:
                conn.commit()
                messagebox.showinfo("Успех", 
                                f"Создано {success_count} бронирований номера {room_number} на сумму {total:.2f} руб.")
                
                # Обновление данных
                self.selected_clients = []
                self.update_selected_clients_list()
                self.update_rooms_table()
                self.fetch_reservations()
                self.display_reservations()
                
                # Обновление статуса номера на "Занят"
                cursor.execute("UPDATE Rooms SET status='Занят' WHERE id=?", (room_id,))
                conn.commit()
                
            else:
                messagebox.showwarning("Предупреждение", "Не удалось создать ни одного бронирования")
                
            conn.close()
            
        except ValueError:
            messagebox.showerror("Ошибка", "Некорректный формат даты (используйте ДД.ММ.ГГГГ)")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось создать бронирование: {str(e)}")
            if conn:
                conn.close()

    def fetch_reservations(self):
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute("""
            SELECT r.id, rm.room_number, c.name, r.checkin_date, r.checkout_date, r.total_price
            FROM Reservations r
            JOIN Rooms rm ON r.room_id = rm.id
            JOIN Clients c ON r.client_id = c.id
            ORDER BY r.room_id, r.checkin_date, r.checkout_date
        """)
        self.reservations = cursor.fetchall()
        conn.close()

    def display_reservations(self):
        for row in self.reservations_tree.get_children():
            self.reservations_tree.delete(row)
        
        # Группировка брони по номерам и датам
        reservations_grouped = {}
        for res in self.reservations:
            room_id, room_number, client_name, checkin_date, checkout_date, total_price = res
            key = (room_id, room_number, checkin_date, checkout_date, total_price)
            
            if key not in reservations_grouped:
                reservations_grouped[key] = []
            reservations_grouped[key].append(client_name)
        
        # Отображение сгруппированных данных
        for key, clients in reservations_grouped.items():
            room_id, room_number, checkin_date, checkout_date, total_price = key
            clients_str = ", ".join(clients)
            
            # Определение статуса брони
            today = datetime.now().date()
            checkin = datetime.strptime(checkin_date, "%Y-%m-%d").date()
            checkout = datetime.strptime(checkout_date, "%Y-%m-%d").date()
            
            if checkout < today:
                status = "Завершено"
            elif checkin > today:
                status = "Предстоящее"
            else:
                status = "Активное"
            
            self.reservations_tree.insert("", "end", 
                                        values=(room_id, room_number, clients_str, 
                                               checkin_date, checkout_date, f"{total_price:.2f}", status))

    def delete_reservation(self):
        if not self.reservations_tree.selection():
            messagebox.showwarning("Предупреждение", "Выберите бронь для удаления")
            return

        selected_reservation = self.reservations_tree.item(self.reservations_tree.selection()[0])['values']
        reservation_id = selected_reservation[0]

        if messagebox.askyesno("Подтверждение", 
                             f"Вы действительно хотите удалить бронь номера {selected_reservation[1]}?"):
            conn = None
            try:
                conn = sqlite3.connect(get_db_path())
                cursor = conn.cursor()
                
                # Удаление бронирования
                cursor.execute("DELETE FROM Reservations WHERE id=?", (reservation_id,))
                
                conn.commit()
                messagebox.showinfo("Успех", "Бронь успешно удалена")
                self.fetch_reservations()
                self.display_reservations()
                self.update_rooms_table()  # Обновление статуса номеров
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось удалить бронь: {str(e)}")
            finally:
                if conn:
                    conn.close()


    def refresh_data(self):
        self.fetch_clients()
        self.fetch_rooms()
        self.update_rooms_table()
        self.fetch_reservations()
        self.display_reservations()


# --- класс, описывающий создание фрейма НОМЕРНОЙ ФОНД ---
class Room(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        testLbl = ttk.Label(self, background='#35A7FF')
        testLbl.pack(expand=True, fill='both')
        self.place(relx=0.3, y=0, relwidth=0.7, relheight=1)
        self.create_widgets()
        self.fetch_data()
        self.display_data()
        self.selected_item = None

    def create_widgets(self):
        # cоздаем notebook для вкладок
        self.notebook = ttk.Notebook(self)
        self.notebook.place(relx=0.05, rely=0.05, relwidth=0.9, relheight=0.9)
        
        # Вкладка для всех номеров
        self.rooms_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.rooms_tab, text="Все номера")
        self.create_rooms_tab()
        
        # Вкладка для номеров по статусам
        self.status_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.status_tab, text="По статусам")
        self.create_status_tab()

    def create_rooms_tab(self):
        # Фрейм для поиска
        search_frame = ttk.Frame(self.rooms_tab)
        search_frame.pack(fill='x', padx=5, pady=5)
        
        # Поля для поиска
        ttk.Label(search_frame, text="Поиск:").pack(side='left', padx=5)
        
        self.search_field = ttk.Combobox(search_frame, values=["Номер", "Класс", "Корпус", "Этаж"])
        self.search_field.pack(side='left', padx=5)
        self.search_field.current(0)
        
        self.search_entry = ttk.Entry(search_frame)
        self.search_entry.pack(side='left', padx=5, expand=True, fill='x')
        
        ttk.Button(search_frame, text="Найти", command=self.search_rooms).pack(side='left', padx=5)
        ttk.Button(search_frame, text="Сброс", command=self.reset_search).pack(side='left', padx=5)
        
        # Кнопки управления
        button_frame = ttk.Frame(self.rooms_tab)
        button_frame.pack(fill='x', padx=5, pady=5)
        
        buttons = [
            ('Новый номер', self.open_add_room_window),
            ('Изменить', self.open_edit_room_window),
            ('Удалить', self.delete_room),
            ('Изменить статус', self.open_change_status_window),
            ('Обновить', self.refresh_data)
        ]
        
        for i, (text, command) in enumerate(buttons):
            btn = ttk.Button(button_frame, text=text, command=command)
            btn.grid(row=0, column=i, padx=2, sticky='ew')
            button_frame.columnconfigure(i, weight=1)
        
        # Таблица номеров
        self.columns = ("ID", "Номер", "Мест", "Класс", "Цена", "Этаж", "Корпус", "Опции", "Статус")
        self.tree = ttk.Treeview(self.rooms_tab, columns=self.columns, show="headings")
        
        for col in self.columns:
            self.tree.heading(col, text=col)
            self.tree.column(col, anchor='c')
            
        # Настройка ширины колонок
        self.tree.column("#1", width=50, stretch=False)
        self.tree.column("#2", width=80, stretch=True)
        self.tree.column("#3", width=60, stretch=True)
        self.tree.column("#4", width=100, stretch=True)
        self.tree.column("#5", width=80, stretch=True)
        self.tree.column("#6", width=60, stretch=True)
        self.tree.column("#7", width=100, stretch=True)
        self.tree.column("#8", width=150, stretch=True)
        self.tree.column("#9", width=120, stretch=True)
        
        self.tree.pack(expand=True, fill='both', padx=5, pady=5)
        self.tree.bind('<<TreeviewSelect>>', self.on_tree_select)

    def create_status_tab(self):
        # Создание фреймов для каждого статуса
        self.status_frames = {}
        statuses = ["Свободен", "Требуется клининг", "Требуется ремонт", "Занят"]
        
        for status in statuses:
            frame = ttk.Frame(self.status_tab)
            self.status_frames[status] = frame
            frame.pack(fill='both', expand=True, padx=5, pady=5)
            
            # Заголовок
            ttk.Label(frame, text=f"Номера со статусом: {status}", font=('Arial', 10, 'bold')).pack()
            
            # Таблица номеров
            tree = ttk.Treeview(frame, columns=("ID", "Номер", "Класс", "Корпус"), show="headings")
            for col in ("ID", "Номер", "Класс", "Корпус"):
                tree.heading(col, text=col)
                tree.column(col, anchor='c')
            
            tree.pack(expand=True, fill='both', padx=5, pady=5)
            self.status_frames[status + "_tree"] = tree

    def on_tree_select(self, event):
        self.selected_item = self.tree.selection()[0] if self.tree.selection() else None

    def fetch_data(self):
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute("""
            SELECT r.id, r.room_number, r.places, rc.class_name, r.price, r.floor, 
                   b.building_name, GROUP_CONCAT(ro.option_name, ', '), r.status
            FROM Rooms r
            LEFT JOIN RoomClasses rc ON r.class_id = rc.id
            LEFT JOIN Buildings b ON r.building_id = b.id
            LEFT JOIN RoomOptions ro ON ro.room_id = r.id
            GROUP BY r.id
        """)
        self.rows = cursor.fetchall()
        conn.close()

    def display_data(self):
        for row in self.tree.get_children():
            self.tree.delete(row)
        
        for row in self.status_frames.values():
            if isinstance(row, ttk.Treeview):
                for item in row.get_children():
                    row.delete(item)
        
        status_trees = {
            "Свободен": self.status_frames["Свободен_tree"],
            "Требуется клининг": self.status_frames["Требуется клининг_tree"],
            "Требуется ремонт": self.status_frames["Требуется ремонт_tree"],
            "Занят": self.status_frames["Занят_tree"]
        }
        
        for room in self.rows:
            room_id, room_number, places, class_name, price, floor, building_name, options, status = room
            self.tree.insert("", "end", values=room)
            
            # Добавление номера в соответствующую таблицу статусов
            if status in status_trees:
                status_trees[status].insert("", "end", values=(
                    room_id, room_number, class_name, building_name
                ))

    def refresh_data(self):
        self.fetch_data()
        self.display_data()

    def search_rooms(self):
        search_field = self.search_field.get()
        search_text = self.search_entry.get()
        
        if not search_text:
            self.refresh_data()
            return
            
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        
        if search_field == "Номер":
            cursor.execute("""
                SELECT r.id, r.room_number, r.places, rc.class_name, r.price, r.floor, 
                       b.building_name, GROUP_CONCAT(ro.option_name, ', '), r.status
                FROM Rooms r
                LEFT JOIN RoomClasses rc ON r.class_id = rc.id
                LEFT JOIN Buildings b ON r.building_id = b.id
                LEFT JOIN RoomOptions ro ON ro.room_id = r.id
                WHERE r.room_number LIKE ?
                GROUP BY r.id
            """, (f"%{search_text}%",))
        elif search_field == "Класс":
            cursor.execute("""
                SELECT r.id, r.room_number, r.places, rc.class_name, r.price, r.floor, 
                       b.building_name, GROUP_CONCAT(ro.option_name, ', '), r.status
                FROM Rooms r
                LEFT JOIN RoomClasses rc ON r.class_id = rc.id
                LEFT JOIN Buildings b ON r.building_id = b.id
                LEFT JOIN RoomOptions ro ON ro.room_id = r.id
                WHERE rc.class_name LIKE ?
                GROUP BY r.id
            """, (f"%{search_text}%",))
        elif search_field == "Корпус":
            cursor.execute("""
                SELECT r.id, r.room_number, r.places, rc.class_name, r.price, r.floor, 
                       b.building_name, GROUP_CONCAT(ro.option_name, ', '), r.status
                FROM Rooms r
                LEFT JOIN RoomClasses rc ON r.class_id = rc.id
                LEFT JOIN Buildings b ON r.building_id = b.id
                LEFT JOIN RoomOptions ro ON ro.room_id = r.id
                WHERE b.building_name LIKE ?
                GROUP BY r.id
            """, (f"%{search_text}%",))
        elif search_field == "Этаж":
            cursor.execute("""
                SELECT r.id, r.room_number, r.places, rc.class_name, r.price, r.floor, 
                       b.building_name, GROUP_CONCAT(ro.option_name, ', '), r.status
                FROM Rooms r
                LEFT JOIN RoomClasses rc ON r.class_id = rc.id
                LEFT JOIN Buildings b ON r.building_id = b.id
                LEFT JOIN RoomOptions ro ON ro.room_id = r.id
                WHERE r.floor = ?
                GROUP BY r.id
            """, (search_text,))
            
        self.rows = cursor.fetchall()
        conn.close()
        self.display_data()

    def reset_search(self):
        self.search_entry.delete(0, tk.END)
        self.refresh_data()

    def open_add_room_window(self):
        self.add_window = tk.Toplevel(self)
        self.add_window.title("Добавить новый номер")
        self.add_window.geometry("500x500")
        
        # Получение списка классов и корпусов
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute("SELECT id, class_name FROM RoomClasses")
        self.classes = cursor.fetchall()
        cursor.execute("SELECT id, building_name FROM Buildings")
        self.buildings = cursor.fetchall()
        cursor.execute("SELECT DISTINCT option_name FROM RoomOptionsList")
        self.options_list = [row[0] for row in cursor.fetchall()]
        conn.close()
        
        # Поля формы
        main_frame = ttk.Frame(self.add_window, padding=10)
        main_frame.pack(fill='both', expand=True)
        
        ttk.Label(main_frame, text="Номер:").grid(row=0, column=0, padx=5, pady=5, sticky='e')
        self.room_number_entry = ttk.Entry(main_frame)
        self.room_number_entry.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(main_frame, text="Количество мест:").grid(row=1, column=0, padx=5, pady=5, sticky='e')
        self.places_entry = ttk.Entry(main_frame)
        self.places_entry.grid(row=1, column=1, padx=5, pady=5)
        
        ttk.Label(main_frame, text="Класс:").grid(row=2, column=0, padx=5, pady=5, sticky='e')
        self.class_combobox = ttk.Combobox(main_frame, values=[c[1] for c in self.classes])
        self.class_combobox.grid(row=2, column=1, padx=5, pady=5)
        
        ttk.Label(main_frame, text="Цена:").grid(row=3, column=0, padx=5, pady=5, sticky='e')
        self.price_entry = ttk.Entry(main_frame)
        self.price_entry.grid(row=3, column=1, padx=5, pady=5)
        
        ttk.Label(main_frame, text="Этаж:").grid(row=4, column=0, padx=5, pady=5, sticky='e')
        self.floor_entry = ttk.Entry(main_frame)
        self.floor_entry.grid(row=4, column=1, padx=5, pady=5)
        
        ttk.Label(main_frame, text="Корпус:").grid(row=5, column=0, padx=5, pady=5, sticky='e')
        self.building_combobox = ttk.Combobox(main_frame, values=[b[1] for b in self.buildings])
        self.building_combobox.grid(row=5, column=1, padx=5, pady=5)
        
        ttk.Label(main_frame, text="Статус:").grid(row=6, column=0, padx=5, pady=5, sticky='e')
        self.status_combobox = ttk.Combobox(main_frame, values=["Свободен", "Требуется клининг", "Требуется ремонт"])
        self.status_combobox.grid(row=6, column=1, padx=5, pady=5)
        self.status_combobox.current(0)
        
        ttk.Label(main_frame, text="Опции:").grid(row=7, column=0, padx=5, pady=5, sticky='ne')
        self.options_frame = ttk.Frame(main_frame)
        self.options_frame.grid(row=7, column=1, padx=5, pady=5, sticky='w')
        
        self.options_vars = {}
        for i, option in enumerate(self.options_list):
            var = tk.IntVar()
            cb = ttk.Checkbutton(self.options_frame, text=option, variable=var)
            cb.grid(row=i, column=0, sticky='w')
            self.options_vars[option] = var
        
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=8, column=0, columnspan=2, pady=10)
        
        ttk.Button(button_frame, text="Добавить", command=self.add_room).pack(side='right', padx=5)
        ttk.Button(button_frame, text="Отмена", command=self.add_window.destroy).pack(side='right', padx=5)
        
        main_frame.columnconfigure(1, weight=1)

    def add_room(self):
        # Получение данных из формы
        room_number = self.room_number_entry.get()
        places = self.places_entry.get()
        class_name = self.class_combobox.get()
        price = self.price_entry.get()
        floor = self.floor_entry.get()
        building_name = self.building_combobox.get()
        status = self.status_combobox.get()
        
        # Валидация
        if not all([room_number, places, class_name, price, floor, building_name]):
            messagebox.showerror("Ошибка", "Все основные поля должны быть заполнены")
            return
        
        try:
            places = int(places)
            price = float(price)
            floor = int(floor)
        except ValueError:
            messagebox.showerror("Ошибка", "Количество мест, цена и этаж должны быть числами")
            return
        
        conn = None
        try:
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            
            # Получение ID класса и корпуса
            cursor.execute("SELECT id FROM RoomClasses WHERE class_name=?", (class_name,))
            class_id = cursor.fetchone()[0]
            
            cursor.execute("SELECT id FROM Buildings WHERE building_name=?", (building_name,))
            building_id = cursor.fetchone()[0]
            
            # Добавление номера
            cursor.execute("""
                INSERT INTO Rooms (room_number, places, class_id, price, floor, building_id, status) 
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (room_number, places, class_id, price, floor, building_id, status))
            room_id = cursor.lastrowid
            
            # Добавляем выбранные опции
            for option, var in self.options_vars.items():
                if var.get() == 1:
                    cursor.execute("""
                        INSERT INTO RoomOptions (room_id, option_name) 
                        VALUES (?, ?)
                    """, (room_id, option))
            
            conn.commit()
            conn.close()
            self.refresh_data()
            self.add_window.destroy()
            messagebox.showinfo("Успех", "Номер успешно добавлен")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось добавить номер: {str(e)}")
        finally:
            if conn:
                conn.close()

    def open_edit_room_window(self):
        if not self.selected_item:
            messagebox.showwarning("Предупреждение", "Выберите номер для редактирования")
            return
            
        # Получаение данных выбранного номера
        room_data = self.tree.item(self.selected_item)['values']
        room_id = room_data[0]
        
        # Получение полных данных о номере из БД
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        
        # Основные данные номера
        cursor.execute("""
            SELECT r.id, r.room_number, r.places, r.price, r.floor, r.status,
                rc.id, rc.class_name, b.id, b.building_name
            FROM Rooms r
            JOIN RoomClasses rc ON r.class_id = rc.id
            JOIN Buildings b ON r.building_id = b.id
            WHERE r.id = ?
        """, (room_id,))
        room_info = cursor.fetchone()
        
        # Опции номера
        cursor.execute("SELECT option_name FROM RoomOptions WHERE room_id = ?", (room_id,))
        room_options = [row[0] for row in cursor.fetchall()]
        
        # Списки для комбобоксов
        cursor.execute("SELECT id, class_name FROM RoomClasses")
        classes = cursor.fetchall()
        
        cursor.execute("SELECT id, building_name FROM Buildings")
        buildings = cursor.fetchall()
        
        cursor.execute("SELECT DISTINCT option_name FROM RoomOptionsList")
        options_list = [row[0] for row in cursor.fetchall()]
        
        conn.close()
        
        # окно редактирования
        self.edit_window = tk.Toplevel(self)
        self.edit_window.title("Редактировать номер")
        self.edit_window.geometry("500x500")
        
        # Поля формы
        main_frame = ttk.Frame(self.edit_window, padding=10)
        main_frame.pack(fill='both', expand=True)
        
        ttk.Label(main_frame, text="Номер:").grid(row=0, column=0, padx=5, pady=5, sticky='e')
        self.edit_room_number = ttk.Entry(main_frame)
        self.edit_room_number.grid(row=0, column=1, padx=5, pady=5)
        self.edit_room_number.insert(0, room_info[1])
        
        ttk.Label(main_frame, text="Количество мест:").grid(row=1, column=0, padx=5, pady=5, sticky='e')
        self.edit_places = ttk.Entry(main_frame)
        self.edit_places.grid(row=1, column=1, padx=5, pady=5)
        self.edit_places.insert(0, room_info[2])
        
        ttk.Label(main_frame, text="Класс:").grid(row=2, column=0, padx=5, pady=5, sticky='e')
        self.edit_class = ttk.Combobox(main_frame, values=[c[1] for c in classes])
        self.edit_class.grid(row=2, column=1, padx=5, pady=5)
        self.edit_class.set(room_info[7])
        
        ttk.Label(main_frame, text="Цена:").grid(row=3, column=0, padx=5, pady=5, sticky='e')
        self.edit_price = ttk.Entry(main_frame)
        self.edit_price.grid(row=3, column=1, padx=5, pady=5)
        self.edit_price.insert(0, room_info[3])
        
        ttk.Label(main_frame, text="Этаж:").grid(row=4, column=0, padx=5, pady=5, sticky='e')
        self.edit_floor = ttk.Entry(main_frame)
        self.edit_floor.grid(row=4, column=1, padx=5, pady=5)
        self.edit_floor.insert(0, room_info[4])
        
        ttk.Label(main_frame, text="Корпус:").grid(row=5, column=0, padx=5, pady=5, sticky='e')
        self.edit_building = ttk.Combobox(main_frame, values=[b[1] for b in buildings])
        self.edit_building.grid(row=5, column=1, padx=5, pady=5)
        self.edit_building.set(room_info[9])
        
        ttk.Label(main_frame, text="Статус:").grid(row=6, column=0, padx=5, pady=5, sticky='e')
        self.edit_status = ttk.Combobox(main_frame, values=["Свободен", "Требуется клининг", "Требуется ремонт", "Занят"])
        self.edit_status.grid(row=6, column=1, padx=5, pady=5)
        self.edit_status.set(room_info[5])
        
        ttk.Label(main_frame, text="Опции:").grid(row=7, column=0, padx=5, pady=5, sticky='ne')
        self.edit_options_frame = ttk.Frame(main_frame)
        self.edit_options_frame.grid(row=7, column=1, padx=5, pady=5, sticky='w')
        
        self.edit_options_vars = {}
        for i, option in enumerate(options_list):
            var = tk.IntVar(value=1 if option in room_options else 0)
            cb = ttk.Checkbutton(self.edit_options_frame, text=option, variable=var)
            cb.grid(row=i, column=0, sticky='w')
            self.edit_options_vars[option] = var
        
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=8, column=0, columnspan=2, pady=10)
        
        ttk.Button(button_frame, text="Сохранить", command=self.edit_room).pack(side='right', padx=5)
        ttk.Button(button_frame, text="Отмена", command=self.edit_window.destroy).pack(side='right', padx=5)
        
        main_frame.columnconfigure(1, weight=1)
        self.editing_room_id = room_id

    def edit_room(self):
        # Получение данных из формы
        room_number = self.edit_room_number.get()
        places = self.edit_places.get()
        class_name = self.edit_class.get()
        price = self.edit_price.get()
        floor = self.edit_floor.get()
        building_name = self.edit_building.get()
        status = self.edit_status.get()
        
        # Валидация данных
        if not all([room_number, places, class_name, price, floor, building_name, status]):
            messagebox.showerror("Ошибка", "Все основные поля должны быть заполнены")
            return
            
        try:
            places = int(places)
            price = float(price)
            floor = int(floor)
        except ValueError:
            messagebox.showerror("Ошибка", "Количество мест, цена и этаж должны быть числами")
            return
        
        conn = None
        try:
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            
            # Получение ID класса и корпуса
            cursor.execute("SELECT id FROM RoomClasses WHERE class_name=?", (class_name,))
            class_id = cursor.fetchone()[0]
            
            cursor.execute("SELECT id FROM Buildings WHERE building_name=?", (building_name,))
            building_id = cursor.fetchone()[0]
            
            # Обновление основные данные номера
            cursor.execute("""
                UPDATE Rooms 
                SET room_number=?, places=?, class_id=?, price=?, floor=?, building_id=?, status=?
                WHERE id=?
            """, (room_number, places, class_id, price, floor, building_id, status, self.editing_room_id))
            
            # Обновление опций номера
            cursor.execute("DELETE FROM RoomOptions WHERE room_id=?", (self.editing_room_id,))
            
            # добавление выбранных опций
            for option, var in self.edit_options_vars.items():
                if var.get() == 1:
                    cursor.execute("""
                        INSERT INTO RoomOptions (room_id, option_name)
                        VALUES (?, ?)
                    """, (self.editing_room_id, option))
            
            conn.commit()
            
            self.refresh_data()
            self.edit_window.destroy()
            messagebox.showinfo("Успех", "Данные номера успешно обновлены")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось обновить данные номера: {str(e)}")
        finally:
            if conn:
                conn.close()

    def delete_room(self):
        if not self.selected_item:
            messagebox.showwarning("Предупреждение", "Выберите номер для удаления")
            return
            
        room_data = self.tree.item(self.selected_item)['values']
        room_id = room_data[0]
        room_number = room_data[1]
        
        # проверка есть ли активные бронирования для этого номера
        try:
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            cursor.execute("""
                SELECT 1 FROM Reservations 
                WHERE room_id=? AND checkout_date >= date('now')
            """, (room_id,))
            
            if cursor.fetchone():
                messagebox.showerror("Ошибка", "Нельзя удалить номер с активными бронированиями")
                conn.close()
                return
                
            if messagebox.askyesno("Подтверждение", 
                                f"Вы действительно хотите удалить номер {room_number}?\nЭто действие нельзя отменить."):
                cursor.execute("DELETE FROM RoomOptions WHERE room_id=?", (room_id,))
                cursor.execute("DELETE FROM Rooms WHERE id=?", (room_id,))
                
                conn.commit()
                conn.close()
                
                self.refresh_data()
                messagebox.showinfo("Успех", "Номер успешно удален")
            else:
                conn.close()
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось удалить номер: {str(e)}")
        finally:
            if conn:
                conn.close()

    def open_change_status_window(self):
        if not self.selected_item:
            messagebox.showwarning("Предупреждение", "Выберите номер для изменения статуса")
            return
            
        room_data = self.tree.item(self.selected_item)['values']
        room_id = room_data[0]
        room_number = room_data[1]
        current_status = room_data[8]
        
        self.status_window = tk.Toplevel(self)
        self.status_window.title(f"Изменить статус номера {room_number}")
        self.status_window.geometry("300x200")
        
        ttk.Label(self.status_window, text=f"Текущий статус: {current_status}").pack(pady=10)
        
        ttk.Label(self.status_window, text="Новый статус:").pack()
        self.new_status = ttk.Combobox(self.status_window, 
                                      values=["Свободен", "Требуется клининг", "Требуется ремонт"])
        self.new_status.pack(pady=10)
        self.new_status.current(0)
        
        ttk.Button(self.status_window, text="Сохранить", 
                  command=lambda: self.change_room_status(room_id)).pack(pady=10)

    def change_room_status(self, room_id):
        new_status = self.new_status.get()
        
        if not new_status:
            messagebox.showerror("Ошибка", "Выберите новый статус")
            return
            
        try:
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            cursor.execute("UPDATE Rooms SET status=? WHERE id=?", (new_status, room_id))
            conn.commit()
            conn.close()
            
            self.status_window.destroy()
            self.refresh_data()
            messagebox.showinfo("Успех", "Статус номера успешно изменен")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось изменить статус: {str(e)}")

# --- Класс, описывающий создание меню в левой части окна ---
class Menu(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        self.configure(style='Menu.TFrame')
        testLbl = ttk.Label(self, background='#38618C')
        testLbl.pack(expand=True, fill='both')
        self.place(x=0, y=0, relwidth=0.3, relheight=1)
        self.create_widgets()
        self.buttonAdmin = ttk.Button(self, text='Администрирование', command=self.show_admin_panel)
        self.update_time()
        self.logged_in = False
        

    def create_widgets(self):
        
        style = ttk.Style()
        style.configure('Menu.TFrame', background='#38618C')
        
        self.buttonMenu1 = ttk.Button(self, text='Клиенты', command=self.show_clients)
        self.buttonMenu2 = ttk.Button(self, text='Бронирование', command=self.show_rent)
        self.buttonMenu3 = ttk.Button(self, text='Номерной фонд', command=self.show_room)
        self.buttonMenu4 = ttk.Button(self, text='Справочники', command = self.show_spravka)
        self.buttonPayment = ttk.Button(self, text='Оплата', command=self.show_payment)
        self.buttonMenu5 = ttk.Button(self, text='О программе', command=self.show_about)
        self.buttonMenu5.place(relx=0.27, rely=0.9, relwidth=0.45, height=40)
        self.buttonMenu6 = ttk.Button(self, text='Вход', command=self.show_login)
        self.buttonMenu6.place(relx=0.27, rely=0.8, relwidth=0.45, height=40)
        self.buttonMenu7 = ttk.Button(self, text='Отчеты', command=self.show_report)
        self.buttonLogout = ttk.Button(self, text='Выход', command=self.logout, state='disabled')


       # Фрейм для даты и времени
        datetime_frame = tk.Frame(self, bg='#38618C', bd=0)
        datetime_frame.place(relx=0.1, rely=0.05, relwidth=0.8, height=80)
        
        # Дата (день недели, число, месяц, год)
        self.date_label = tk.Label(
            datetime_frame,
            bg='#38618C',
            fg='white',
            font=('Arial', 12, 'bold'),
            anchor='center'
        )
        self.date_label.pack(fill='x', pady=(5, 0))
        
        # Время (часы:минуты:секунды)
        self.time_label = tk.Label(
            datetime_frame,
            bg='#38618C',
            fg='yellow',
            font=('Arial', 16, 'bold'),
            anchor='center'
        )
        self.time_label.pack(fill='x', pady=(0, 5))

    # Обновление отображения даты и времени
    def update_time(self):
        now = datetime.now()
        
        # Форматирование даты: "05.06.2025" (день.месяц.год)
        current_date = now.strftime("%d.%m.%Y")
        
        # Форматирование времени: "14:30:45" (оставляем как было)
        current_time = now.strftime("%H:%M:%S")
        
        # Обновление метки
        self.date_label.config(text=current_date)
        self.time_label.config(text=current_time)
        
        # вызов этой функции снова через 1000 мс (1 секунду)
        self.after(1000, self.update_time)

    # Сокрытие основных кнопки меню
    def hide_main_buttons(self):
        self.buttonMenu1.place_forget()
        self.buttonMenu2.place_forget()
        self.buttonMenu3.place_forget()
        self.buttonMenu4.place_forget()
        self.buttonMenu7.place_forget()
        self.buttonPayment.place_forget()
        self.buttonLogout.place_forget()
        self.logged_in = False
        self.buttonLogout.config(state='disabled')

    # Отображение основных кнопок меню после успешного входа
    def show_main_buttons(self):
        self.buttonMenu1.place(relx=0.27, rely=0.2, relwidth=0.45, height=40)
        self.buttonMenu2.place(relx=0.27, rely=0.3, relwidth=0.45, height=40)
        self.buttonMenu3.place(relx=0.27, rely=0.4, relwidth=0.45, height=40)
        self.buttonMenu4.place(relx=0.27, rely=0.5, relwidth=0.45, height=40)
        self.buttonMenu7.place(relx=0.27, rely=0.6, relwidth=0.45, height=40)
        self.buttonPayment.place(relx=0.27, rely=0.7, relwidth=0.45, height=40)
        self.buttonMenu6.place_forget()
        self.buttonLogout.place(relx=0.27, rely=0.8, relwidth=0.45, height=40)
        self.buttonLogout.config(state='normal')
        self.logged_in = True
    
    # Показывает кнопку администрирования для суперпользователя
    def show_admin_button(self):
        
        self.buttonAdmin.place(relx=0.27, rely=0.15, relwidth=0.45, height=40)
    
    # Открывает панель администрирования
    def show_admin_panel(self):
        
        admin_window = tk.Toplevel(self)
        admin_window.title("Администрирование")
        admin_window.geometry("600x300")
        
        # Таблица пользователей
        tree = ttk.Treeview(admin_window, columns=("ID", "Логин"), show="headings")
        tree.heading("ID", text="ID")
        tree.heading("Логин", text="Логин")
        tree.pack(expand=True, fill='both', padx=5, pady=5)
        
        # Загрузка пользователей
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute("SELECT id, name FROM Users")
        users = cursor.fetchall()
        conn.close()
        
        for user in users:
            tree.insert("", "end", values=user)
            
        # Кнопки управления
        button_frame = ttk.Frame(admin_window)
        button_frame.pack(fill='x', padx=5, pady=5)
        
        ttk.Button(button_frame, text="Добавить", command=lambda: self.add_user(tree)).pack(side='left', padx=5)
        ttk.Button(button_frame, text="Удалить", command=lambda: self.delete_user(tree)).pack(side='left', padx=5)
        ttk.Button(button_frame, text="Изменить пароль", command=lambda: self.change_password(tree)).pack(side='left', padx=5)
        
    # Добавляет нового пользователя
    def add_user(self, tree):
        add_window = tk.Toplevel(self)
        add_window.title("Добавить пользователя")
        add_window.geometry("300x200")
        
        ttk.Label(add_window, text="Логин:").pack(pady=5)
        username_entry = ttk.Entry(add_window)
        username_entry.pack(pady=5)
        
        ttk.Label(add_window, text="Пароль:").pack(pady=5)
        password_entry = ttk.Entry(add_window, show="*")
        password_entry.pack(pady=5)
        
        def save_user():
            username = username_entry.get()
            password = password_entry.get()
            
            if not username or not password:
                messagebox.showerror("Ошибка", "Введите логин и пароль")
                return
                
            try:
                conn = sqlite3.connect(get_db_path())
                cursor = conn.cursor()
                cursor.execute("INSERT INTO Users (name, password) VALUES (?, ?)", (username, password))
                conn.commit()
                conn.close()
                
                # Обновляем таблицу
                tree.insert("", "end", values=(cursor.lastrowid, username))
                add_window.destroy()
                messagebox.showinfo("Успех", "Пользователь добавлен")
            except sqlite3.IntegrityError:
                messagebox.showerror("Ошибка", "Пользователь с таким логином уже существует")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось добавить пользователя: {str(e)}")
                
        ttk.Button(add_window, text="Сохранить", command=save_user).pack(pady=10)
        
    # Удаляет пользователя
    def delete_user(self, tree):
        if not tree.selection():
            messagebox.showwarning("Предупреждение", "Выберите пользователя для удаления")
            return
            
        user_id = tree.item(tree.selection()[0])['values'][0]
        username = tree.item(tree.selection()[0])['values'][1]
        
        # Защищаем учетную запись root от удаления
        if username.lower() == "root":
            messagebox.showerror("Ошибка", "Нельзя удалить суперпользователя root")
            return
            
        if messagebox.askyesno("Подтверждение", f"Удалить пользователя {username}?"):
            try:
                conn = sqlite3.connect(get_db_path())
                cursor = conn.cursor()
                cursor.execute("DELETE FROM Users WHERE id=?", (user_id,))
                conn.commit()
                conn.close()
                
                tree.delete(tree.selection()[0])
                messagebox.showinfo("Успех", "Пользователь удален")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось удалить пользователя: {str(e)}")
                
    # Изменяет пароль пользователя
    def change_password(self, tree):
        if not tree.selection():
            messagebox.showwarning("Предупреждение", "Выберите пользователя")
            return
            
        user_id = tree.item(tree.selection()[0])['values'][0]
        username = tree.item(tree.selection()[0])['values'][1]
        
        change_window = tk.Toplevel(self)
        change_window.title("Изменить пароль")
        change_window.geometry("300x150")
        
        ttk.Label(change_window, text=f"Пользователь: {username}").pack(pady=5)
        ttk.Label(change_window, text="Новый пароль:").pack(pady=5)
        password_entry = ttk.Entry(change_window, show="*")
        password_entry.pack(pady=5)
        
        def save_password():
            password = password_entry.get()
            
            if not password:
                messagebox.showerror("Ошибка", "Введите пароль")
                return
                
            try:
                conn = sqlite3.connect(get_db_path())
                cursor = conn.cursor()
                cursor.execute("UPDATE Users SET password=? WHERE id=?", (password, user_id))
                conn.commit()
                conn.close()
                
                change_window.destroy()
                messagebox.showinfo("Успех", "Пароль изменен")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось изменить пароль: {str(e)}")
                
        ttk.Button(change_window, text="Сохранить", command=save_password).pack(pady=10)

    # Скрывает кнопку администрирования
    def hide_admin_button(self):
        self.buttonAdmin.place_forget()

    # Выход из системы
    def logout(self):
        self.hide_main_buttons()
        self.hide_admin_button()
        self.show_login()
        self.buttonMenu6.place(relx=0.27, rely=0.8, relwidth=0.45, height=40)  # Показываем кнопку входа
        messagebox.showinfo("Выход", "Вы успешно вышли из системы")

    def set_spravka_frame(self, spravka_frame):
        self.spravka_frame = spravka_frame

    def show_spravka(self):
        self.spravka_frame.tkraise()

    def set_payment_frame(self, payment_frame):
        self.payment_frame = payment_frame
    
    def show_payment(self):
        self.payment_frame.tkraise()

    def set_about_frame(self, about_frame):
        self.about_frame = about_frame

    def show_about(self):
        self.about_frame.tkraise()


    def set_clients_frame(self, clients_frame):
        self.clients_frame = clients_frame

    def show_clients(self):
        self.clients_frame.tkraise()


    def set_rent_frame(self, rent_frame):
        self.rent_frame = rent_frame

    def show_rent(self):
        self.rent_frame.tkraise()


    def set_room_frame(self, room_frame):
        self.room_frame = room_frame

    def show_room(self):
        self.room_frame.tkraise()


    def set_login_frame(self, login_frame):
        self.login_frame = login_frame

    def show_login(self):
        self.login_frame.tkraise()

    def set_report_frame(self, report_frame):
        self.report_frame = report_frame
    
    def show_report(self):
        self.report_frame.tkraise()
        
    def set_admin_frame(self, admin_frame):
        self.admin_frame = admin_frame
    
    def show_admin_panel(self):
        self.admin_frame.tkraise()

# --- Класс, описывающий создание фрейма АВТОРИЗАЦИЯ ---
class Login(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        testLbl = ttk.Label(self, background="#35A7FF")
        testLbl.pack(expand=True, fill='both')
        self.place(relx=0.3, y=0, relwidth=0.7, relheight=1)
        self.create_widgets()
        self.bind('<Return>', self.login)

    def create_widgets(self):
        # Стилизация
        style = ttk.Style()
        style.configure('Login.TFrame', background="#942e2e")
        style.configure('Login.TLabel', background="#ab1616", font=('Arial', 10))
        style.configure('Login.TEntry', font=('Arial', 10), padding=5)
        style.configure('Login.TButton', font=('Arial', 10, 'bold'), padding=5)
        
        # Основной контейнер
        container = ttk.Frame(self, style='Login.TFrame')
        #ttk.Label(container, background="#23699E").pack(expand=True, fill='both')
        container.place(relx=0.5, rely=0.5, anchor='center', width=400, height=300)
        
        # Заголовок
        
        ttk.Label(container, text='Авторизация', style='Login.TLabel', 
                 font=('Arial', 14, 'bold')).pack(pady=(20, 30))
        
        
        
        # Поля ввода
        input_frame = ttk.Frame(container, style='Login.TFrame')
        input_frame.pack(pady=10)
        
        ttk.Label(input_frame, text='Логин:', style='Login.TLabel').grid(row=0, column=0, padx=5, pady=5, sticky='e')
        self.entryLoginLogin = ttk.Entry(input_frame, style='Login.TEntry')
        self.entryLoginLogin.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(input_frame, text='Пароль:', style='Login.TLabel').grid(row=1, column=0, padx=5, pady=5, sticky='e')
        self.entryLoginPassword = ttk.Entry(input_frame, show='*', style='Login.TEntry')
        self.entryLoginPassword.grid(row=1, column=1, padx=5, pady=5)
        
        # Кнопка входа
        button_frame = ttk.Frame(container, style='Login.TFrame')
        button_frame.pack(pady=20)
        
        self.buttonLoginEnter = ttk.Button(button_frame, text='Войти', command=self.login, style='Login.TButton')
        self.buttonLoginEnter.pack(pady=5, ipadx=20)
        
        # Метка для статуса входа
        self.login_status = ttk.Label(container, text='', style='Login.TLabel')
        self.login_status.pack()

        # Устанавливаем фокус на поле логина
        self.entryLoginLogin.focus_set()

        

    def login(self, event=None):
        username = self.entryLoginLogin.get()
        password = self.entryLoginPassword.get()
        
        if not username or not password:
            self.login_status.config(text='Введите логин и пароль', foreground='red')
            return
            
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM Users WHERE name = ? AND password = ?;', (username, password))
        user = cursor.fetchone()
        conn.close()

        if user:
            self.login_status.config(text='Успешный вход', foreground='green')
            self.master.menu.show_main_buttons()
            
            # Проверяем, является ли пользователь суперпользователем (root)
            if username.lower() == "root":
                self.master.menu.show_admin_button()
                
            self.master.clients.tkraise()
        else:
            self.login_status.config(text='Ошибка входа: неверный логин или пароль', foreground='red')


# --- Класс, описывающий создание фрейма СПРАВОЧНИКИ ---
class Spravka(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        self.parent = parent
        testLbl = ttk.Label(self, background='#35A7FF')
        testLbl.pack(expand=True, fill='both')
        self.place(relx=0.3, y=0, relwidth=0.7, relheight=1)
        self.create_widgets()
        
    def create_widgets(self):
        # Создаем Notebook для вкладок
        self.notebook = ttk.Notebook(self)
        self.notebook.place(relx=0.05, rely=0.05, relwidth=0.9, relheight=0.9)
        
        # Вкладка для классов номеров
        self.classes_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.classes_tab, text="Классы номеров")
        self.create_classes_tab()
        
        # Вкладка для корпусов
        self.buildings_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.buildings_tab, text="Корпуса")
        self.create_buildings_tab()
        
        # Вкладка для опций номеров
        self.options_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.options_tab, text="Опции номеров")
        self.create_options_tab()
    
    def create_classes_tab(self):
        # Таблица классов
        self.classes_tree = ttk.Treeview(self.classes_tab, columns=("ID", "Класс", "Описание"), show="headings")
        for col in ("ID", "Класс", "Описание"):
            self.classes_tree.heading(col, text=col)
        self.classes_tree.pack(expand=True, fill='both', padx=5, pady=5)

        self.classes_tree.column("#1", stretch=True, width=40, anchor='c')
        self.classes_tree.column("#2", stretch=True, width=120, anchor='c')
        self.classes_tree.column("#3", stretch=True, width=300, anchor='c')
        
        
        # Кнопки управления
        button_frame = ttk.Frame(self.classes_tab)
        button_frame.pack(fill='x', padx=5, pady=5)
        
        ttk.Button(button_frame, text="Добавить", command=self.open_add_class_window).pack(side='left', padx=5)
        ttk.Button(button_frame, text="Изменить", command=self.open_edit_class_window).pack(side='left', padx=5)
        ttk.Button(button_frame, text="Удалить", command=self.delete_class).pack(side='left', padx=5)
        ttk.Button(button_frame, text="Обновить", command=self.refresh_classes).pack(side='right', padx=5)
        
        self.refresh_classes()
    
    def refresh_classes(self):
        
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute("SELECT id, class_name, description FROM RoomClasses")
        self.classes = cursor.fetchall()
        conn.close()
        
        for row in self.classes_tree.get_children():
            self.classes_tree.delete(row)
        
        for cls in self.classes:
            self.classes_tree.insert("", "end", values=cls)
    
    def open_add_class_window(self):
        self.add_class_window = tk.Toplevel(self)
        self.add_class_window.title("Добавить класс номера")
        
        tk.Label(self.add_class_window, text="Название класса:").grid(row=0, column=0, padx=5, pady=5, sticky='e')
        self.class_name_entry = tk.Entry(self.add_class_window)
        self.class_name_entry.grid(row=0, column=1, padx=5, pady=5)
        
        tk.Label(self.add_class_window, text="Описание:").grid(row=1, column=0, padx=5, pady=5, sticky='ne')
        self.class_desc_text = tk.Text(self.add_class_window, width=30, height=5)
        self.class_desc_text.grid(row=1, column=1, padx=5, pady=5)
        
        ttk.Button(self.add_class_window, text="Добавить", command=self.add_class).grid(row=2, column=1, pady=10)
    
    def add_class(self):
        name = self.class_name_entry.get()
        desc = self.class_desc_text.get("1.0", tk.END).strip()
        
        if not name:
            messagebox.showerror("Ошибка", "Название класса обязательно")
            return
            
        try:
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            cursor.execute("INSERT INTO RoomClasses (class_name, description) VALUES (?, ?)", (name, desc))
            conn.commit()
            conn.close()
            
            self.refresh_classes()
            self.add_class_window.destroy()
            messagebox.showinfo("Успех", "Класс успешно добавлен")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось добавить класс: {str(e)}")
    
    def open_edit_class_window(self):
        if not self.classes_tree.selection():
            messagebox.showwarning("Предупреждение", "Выберите класс для редактирования")
            return
        
        class_id = self.classes_tree.item(self.classes_tree.selection()[0])['values'][0]
        
        # Получаем данные выбранного класса
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute("SELECT id, class_name, description FROM RoomClasses WHERE id=?", (class_id,))
        class_data = cursor.fetchone()
        conn.close()
        
        # Создаем окно редактирования
        self.edit_class_window = tk.Toplevel(self)
        self.edit_class_window.title("Редактировать класс номера")
        
        tk.Label(self.edit_class_window, text="Название класса:").grid(row=0, column=0, padx=5, pady=5, sticky='e')
        self.edit_class_name = tk.Entry(self.edit_class_window)
        self.edit_class_name.grid(row=0, column=1, padx=5, pady=5)
        self.edit_class_name.insert(0, class_data[1])
        
        tk.Label(self.edit_class_window, text="Описание:").grid(row=1, column=0, padx=5, pady=5, sticky='ne')
        self.edit_class_desc = tk.Text(self.edit_class_window, width=30, height=5)
        self.edit_class_desc.grid(row=1, column=1, padx=5, pady=5)
        self.edit_class_desc.insert("1.0", class_data[2])
        
        self.editing_class_id = class_id
        
        ttk.Button(self.edit_class_window, text="Сохранить", command=self.edit_class).grid(row=2, column=1, pady=10)

    def edit_class(self):
        name = self.edit_class_name.get()
        desc = self.edit_class_desc.get("1.0", tk.END).strip()
        
        if not name:
            messagebox.showerror("Ошибка", "Название класса обязательно")
            return
            
        try:
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            cursor.execute("""
                UPDATE RoomClasses 
                SET class_name=?, description=? 
                WHERE id=?
            """, (name, desc, self.editing_class_id))
            conn.commit()
            conn.close()
            
            self.refresh_classes()
            self.edit_class_window.destroy()
            messagebox.showinfo("Успех", "Данные класса успешно обновлены")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось обновить данные класса: {str(e)}")

    def delete_class(self):
        if not self.classes_tree.selection():
            messagebox.showwarning("Предупреждение", "Выберите класс для удаления")
            return
        
        class_id = self.classes_tree.item(self.classes_tree.selection()[0])['values'][0]
        class_name = self.classes_tree.item(self.classes_tree.selection()[0])['values'][1]
        
        # Проверяем, есть ли номера этого класса
        try:
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            cursor.execute("SELECT 1 FROM Rooms WHERE class_id=?", (class_id,))
            
            if cursor.fetchone():
                messagebox.showerror("Ошибка", "Нельзя удалить класс, к которому привязаны номера")
                conn.close()
                return
                
            if messagebox.askyesno("Подтверждение", 
                                f"Вы действительно хотите удалить класс {class_name}?\nЭто действие нельзя отменить."):
                cursor.execute("DELETE FROM RoomClasses WHERE id=?", (class_id,))
                conn.commit()
                conn.close()
                
                self.refresh_classes()
                messagebox.showinfo("Успех", "Класс успешно удален")
            else:
                conn.close()
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось удалить класс: {str(e)}")

    def create_buildings_tab(self):
        # Таблица корпусов
        self.buildings_tree = ttk.Treeview(self.buildings_tab, columns=("ID", "Корпус", "Описание"), show="headings")
        for col in ("ID", "Корпус", "Описание"):
            self.buildings_tree.heading(col, text=col)
        self.buildings_tree.pack(expand=True, fill='both', padx=5, pady=5)

        self.buildings_tree.column("#1", stretch=True, width=40, anchor='c')
        self.buildings_tree.column("#2", stretch=True, width=120, anchor='c')
        self.buildings_tree.column("#3", stretch=True, width=300, anchor='c')     
        
        # Кнопки управления
        button_frame = ttk.Frame(self.buildings_tab)
        button_frame.pack(fill='x', padx=5, pady=5)
        
        ttk.Button(button_frame, text="Добавить", command=self.open_add_building_window).pack(side='left', padx=5)
        ttk.Button(button_frame, text="Изменить", command=self.open_edit_building_window).pack(side='left', padx=5)
        ttk.Button(button_frame, text="Удалить", command=self.delete_building).pack(side='left', padx=5)
        ttk.Button(button_frame, text="Обновить", command=self.refresh_buildings).pack(side='right', padx=5)
        
        self.refresh_buildings()

    def refresh_buildings(self):
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute("SELECT id, building_name, description FROM Buildings")
        self.buildings = cursor.fetchall()
        conn.close()
        
        for row in self.buildings_tree.get_children():
            self.buildings_tree.delete(row)
        
        for building in self.buildings:
            self.buildings_tree.insert("", "end", values=building)

    def open_add_building_window(self):
        self.add_building_window = tk.Toplevel(self)
        self.add_building_window.title("Добавить корпус")
        
        tk.Label(self.add_building_window, text="Название корпуса:").grid(row=0, column=0, padx=5, pady=5, sticky='e')
        self.building_name_entry = tk.Entry(self.add_building_window)
        self.building_name_entry.grid(row=0, column=1, padx=5, pady=5)
        
        tk.Label(self.add_building_window, text="Описание:").grid(row=1, column=0, padx=5, pady=5, sticky='ne')
        self.building_desc_text = tk.Text(self.add_building_window, width=30, height=5)
        self.building_desc_text.grid(row=1, column=1, padx=5, pady=5)
        
        ttk.Button(self.add_building_window, text="Добавить", command=self.add_building).grid(row=2, column=1, pady=10)

    def add_building(self):
        name = self.building_name_entry.get()
        desc = self.building_desc_text.get("1.0", tk.END).strip()
        
        if not name:
            messagebox.showerror("Ошибка", "Название корпуса обязательно")
            return
            
        try:
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            cursor.execute("INSERT INTO Buildings (building_name, description) VALUES (?, ?)", (name, desc))
            conn.commit()
            conn.close()
            
            self.refresh_buildings()
            self.add_building_window.destroy()
            messagebox.showinfo("Успех", "Корпус успешно добавлен")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось добавить корпус: {str(e)}")

    def open_edit_building_window(self):
        if not self.buildings_tree.selection():
            messagebox.showwarning("Предупреждение", "Выберите корпус для редактирования")
            return
        
        building_id = self.buildings_tree.item(self.buildings_tree.selection()[0])['values'][0]
        
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute("SELECT id, building_name, description FROM Buildings WHERE id=?", (building_id,))
        building_data = cursor.fetchone()
        conn.close()
        
        self.edit_building_window = tk.Toplevel(self)
        self.edit_building_window.title("Редактировать корпус")
        
        tk.Label(self.edit_building_window, text="Название корпуса:").grid(row=0, column=0, padx=5, pady=5, sticky='e')
        self.edit_building_name = tk.Entry(self.edit_building_window)
        self.edit_building_name.grid(row=0, column=1, padx=5, pady=5)
        self.edit_building_name.insert(0, building_data[1])
        
        tk.Label(self.edit_building_window, text="Описание:").grid(row=1, column=0, padx=5, pady=5, sticky='ne')
        self.edit_building_desc = tk.Text(self.edit_building_window, width=30, height=5)
        self.edit_building_desc.grid(row=1, column=1, padx=5, pady=5)
        self.edit_building_desc.insert("1.0", building_data[2])
        
        self.editing_building_id = building_id
        
        ttk.Button(self.edit_building_window, text="Сохранить", command=self.edit_building).grid(row=2, column=1, pady=10)

    def edit_building(self):
        name = self.edit_building_name.get()
        desc = self.edit_building_desc.get("1.0", tk.END).strip()
        
        if not name:
            messagebox.showerror("Ошибка", "Название корпуса обязательно")
            return
            
        try:
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            cursor.execute("""
                UPDATE Buildings 
                SET building_name=?, description=? 
                WHERE id=?
            """, (name, desc, self.editing_building_id))
            conn.commit()
            conn.close()
            
            self.refresh_buildings()
            self.edit_building_window.destroy()
            messagebox.showinfo("Успех", "Данные корпуса успешно обновлены")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось обновить данные корпуса: {str(e)}")

    def delete_building(self):
        if not self.buildings_tree.selection():
            messagebox.showwarning("Предупреждение", "Выберите корпус для удаления")
            return
        
        building_id = self.buildings_tree.item(self.buildings_tree.selection()[0])['values'][0]
        building_name = self.buildings_tree.item(self.buildings_tree.selection()[0])['values'][1]
        
        # Проверяем, есть ли номера в этом корпусе
        try:
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            cursor.execute("SELECT 1 FROM Rooms WHERE building_id=?", (building_id,))
            
            if cursor.fetchone():
                messagebox.showerror("Ошибка", "Нельзя удалить корпус, в котором есть номера")
                conn.close()
                return
                
            if messagebox.askyesno("Подтверждение", 
                                f"Вы действительно хотите удалить корпус {building_name}?\nЭто действие нельзя отменить."):
                cursor.execute("DELETE FROM Buildings WHERE id=?", (building_id,))
                conn.commit()
                conn.close()
                
                self.refresh_buildings()
                messagebox.showinfo("Успех", "Корпус успешно удален")
            else:
                conn.close()
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось удалить корпус: {str(e)}")

    def create_options_tab(self):
        # Таблица опций
        self.options_tree = ttk.Treeview(self.options_tab, columns=("ID", "Опция"), show="headings")
        for col in ("ID", "Опция"):
            self.options_tree.heading(col, text=col)
        self.options_tree.pack(expand=True, fill='both', padx=5, pady=5)

        self.options_tree.column("#1", stretch=True, width=40, anchor='c')
        self.options_tree.column("#2", stretch=True, width=200, anchor='c')
        
        
        # Кнопки управления
        button_frame = ttk.Frame(self.options_tab)
        button_frame.pack(fill='x', padx=5, pady=5)
        
        ttk.Button(button_frame, text="Добавить", command=self.open_add_option_window).pack(side='left', padx=5)
        ttk.Button(button_frame, text="Удалить", command=self.delete_option).pack(side='left', padx=5)
        ttk.Button(button_frame, text="Обновить", command=self.refresh_options).pack(side='right', padx=5)
        
        self.refresh_options()

    def refresh_options(self):
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute("SELECT id, option_name FROM RoomOptionsList")
        self.options = cursor.fetchall()
        conn.close()
        
        for row in self.options_tree.get_children():
            self.options_tree.delete(row)
        
        for option in self.options:
            self.options_tree.insert("", "end", values=option)

    def open_add_option_window(self):
        self.add_option_window = tk.Toplevel(self)
        self.add_option_window.title("Добавить опцию")
        self.add_option_window.geometry("300x150")
        
        tk.Label(self.add_option_window, text="Название опции:").pack(pady=5)
        self.option_name_entry = tk.Entry(self.add_option_window)
        self.option_name_entry.pack(pady=5)
        
        ttk.Button(self.add_option_window, text="Добавить", command=self.add_option).pack(pady=10)

    def add_option(self):
        name = self.option_name_entry.get().strip()
        
        if not name:
            messagebox.showerror("Ошибка", "Название опции не может быть пустым")
            return
            
        try:
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            cursor.execute("INSERT INTO RoomOptionsList (option_name) VALUES (?)", (name,))
            conn.commit()
            conn.close()
            
            self.refresh_options()
            self.add_option_window.destroy()
            messagebox.showinfo("Успех", "Опция успешно добавлена")
        except sqlite3.IntegrityError:
            messagebox.showerror("Ошибка", "Опция с таким названием уже существует")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось добавить опцию: {str(e)}")

    def delete_option(self):
        if not self.options_tree.selection():
            messagebox.showwarning("Предупреждение", "Выберите опцию для удаления")
            return
        
        option_id = self.options_tree.item(self.options_tree.selection()[0])['values'][0]
        option_name = self.options_tree.item(self.options_tree.selection()[0])['values'][1]
        
        # проверяем, используется ли опция в каких-либо номерах
        try:
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            cursor.execute("SELECT 1 FROM RoomOptions WHERE option_name=?", (option_name,))
            
            if cursor.fetchone():
                messagebox.showerror("Ошибка", "Нельзя удалить опцию, которая используется в номерах")
                conn.close()
                return
                
            if messagebox.askyesno("Подтверждение", 
                                f"Вы действительно хотите удалить опцию {option_name}?\nЭто действие нельзя отменить."):
                cursor.execute("DELETE FROM RoomOptionsList WHERE id=?", (option_id,))
                conn.commit()
                conn.close()
                
                self.refresh_options()
                messagebox.showinfo("Успех", "Опция успешно удалена")
            else:
                conn.close()
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось удалить опцию: {str(e)}")

# --- Класс, описывающий создание фрейма ОТЧЕТЫ ---
class Report(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        testLbl = ttk.Label(self, background='#35A7FF')
        testLbl.pack(expand=True, fill='both')
        self.place(relx=0.3, y=0, relwidth=0.7, relheight=1)
        self.create_widgets()
        
    def create_widgets(self):
        # Создаем notebook для вкладок
        self.notebook = ttk.Notebook(self)
        self.notebook.place(relx=0.05, rely=0.05, relwidth=0.9, relheight=0.85)
        
        # Вкладка для финансов
        self.finance_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.finance_tab, text="Финансы")
        self.create_finance_tab()
        
        # Кнопки экспорта и очистки
        button_frame = ttk.Frame(self)
        button_frame.place(relx=0.05, rely=0.92, relwidth=0.9)
        
        ttk.Button(button_frame, text="Экспорт в Excel", command=self.export_current_tab_to_excel).pack(side='left', padx=5)
        ttk.Button(button_frame, text="Экспорт в Word", command=self.export_current_tab_to_word).pack(side='left', padx=5)
        ttk.Button(button_frame, text="Очистить финансовые данные", command=self.clear_financial_data).pack(side='right', padx=5)
    
    def create_finance_tab(self):
        # Фрейм для поиска
        search_frame = ttk.Frame(self.finance_tab)
        search_frame.pack(fill='x', padx=5, pady=5)
        
        ttk.Label(search_frame, text="Поиск:").pack(side='left', padx=5)
        
        self.finance_search_field = ttk.Combobox(search_frame, 
                                              values=["Тип", "Сумма", "Дата", "Описание"], 
                                              width=12, state='readonly')
        self.finance_search_field.current(0)
        self.finance_search_field.pack(side='left', padx=5)
        
        self.finance_search_entry = ttk.Entry(search_frame)
        self.finance_search_entry.pack(side='left', padx=5, expand=True, fill='x')
        
        ttk.Button(search_frame, text="Найти", command=self.search_finance).pack(side='left', padx=5)
        ttk.Button(search_frame, text="Сброс", command=self.reset_finance_search).pack(side='left', padx=5)
        
        # Фрейм для управления финансами
        control_frame = ttk.Frame(self.finance_tab)
        control_frame.pack(fill='x', padx=5, pady=5)
        
        # Кнопка добавления расхода
        ttk.Button(control_frame, text="Добавить расход", command=self.open_add_expense_window).pack(side='left', padx=5)
        
        # Кнопка обновления
        ttk.Button(control_frame, text="Обновить", command=self.generate_finance_report).pack(side='right', padx=5)
        
        ttk.Button(control_frame, text="Удалить", command=self.delete_finance_record).pack(side='left', padx=5)
        ttk.Button(control_frame, text="Изменить", command=self.open_edit_finance_window).pack(side='left', padx=5)

        # Таблица для финансового отчета
        self.finance_tree = ttk.Treeview(
            self.finance_tab, 
            columns=("ID", "Тип", "Сумма", "Дата", "Описание"), 
            show="headings"
        )
        
        # Настройка колонок
        columns = {
            "ID": {"width": 50, "anchor": "center"},
            "Тип": {"width": 100, "anchor": "center"},
            "Сумма": {"width": 100, "anchor": "e"},
            "Дата": {"width": 100, "anchor": "center"},
            "Описание": {"width": 200, "anchor": "w"}
        }
        
        for col, params in columns.items():
            self.finance_tree.heading(col, text=col)
            self.finance_tree.column(col, **params)
        
        self.finance_tree.pack(expand=True, fill='both', padx=5, pady=5)
        
        # Подвал с итогами
        footer_frame = ttk.Frame(self.finance_tab)
        footer_frame.pack(fill='x', padx=5, pady=5)
        
        ttk.Label(footer_frame, text="Доходы:").pack(side='left', padx=5)
        self.income_label = ttk.Label(footer_frame, text="0 руб.", font=('Arial', 10, 'bold'))
        self.income_label.pack(side='left', padx=5)
        
        ttk.Label(footer_frame, text="Расходы:").pack(side='left', padx=5)
        self.expense_label = ttk.Label(footer_frame, text="0 руб.", font=('Arial', 10, 'bold'))
        self.expense_label.pack(side='left', padx=5)
        
        ttk.Label(footer_frame, text="Итого:").pack(side='left', padx=5)
        self.total_label = ttk.Label(footer_frame, text="0 руб.", font=('Arial', 10, 'bold'))
        self.total_label.pack(side='left', padx=5)
        
        self.generate_finance_report()
    
    # Поиск финансовых записей
    def search_finance(self):
        search_field = self.finance_search_field.get()
        search_text = self.finance_search_entry.get().strip()
        
        if not search_text:
            self.generate_finance_report()
            return
            
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        
        query = "SELECT id, type, amount, date, description FROM Finances WHERE 1=1"
        
        if search_field == "Тип":
            query += " AND type LIKE ?"
            search_text = f"%{search_text.lower()}%"
        elif search_field == "Сумма":
            try:
                search_num = float(search_text)
                query += " AND amount = ?"
                search_text = search_num
            except ValueError:
                messagebox.showerror("Ошибка", "Введите число для поиска по сумме")
                return
        elif search_field == "Дата":
            query += " AND date LIKE ?"
            search_text = f"%{search_text}%"
        elif search_field == "Описание":
            query += " AND description LIKE ?"
            search_text = f"%{search_text}%"
            
        cursor.execute(query, (search_text,))
        finance_data = cursor.fetchall()
        conn.close()
        
        # Обновляем таблицу
        for row in self.finance_tree.get_children():
            self.finance_tree.delete(row)
        
        for row in finance_data:
            item_id, type_, amount, date, description = row
            type_text = "Расход" if type_ == "expense" else "Доход"
            self.finance_tree.insert(
                "", "end", 
                values=(item_id, type_text, f"{amount:.2f} руб.", date, description)
            )
    
    # Сброс поиска финансовых записей
    def reset_finance_search(self):
        self.finance_search_entry.delete(0, tk.END)
        self.generate_finance_report()
    
    # Добавим метод для очистки финансовых данных
    def clear_financial_data(self):
        if not messagebox.askyesno("Подтверждение", 
                                "Вы действительно хотите очистить все финансовые данные?\nЭто действие нельзя отменить."):
            return
        
        try:
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            
            # Удаляем все финансовые записи
            cursor.execute("DELETE FROM Finances")
            
            # Удаляем все платежи
            cursor.execute("DELETE FROM Payments")
            
            conn.commit()
            conn.close()
            
            # Обновляем отчеты
            self.generate_finance_report()
            self.generate_top_clients_report()
            self.generate_top_rooms_report()
            
            messagebox.showinfo("Успех", "Все финансовые данные были очищены")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось очистить данные: {str(e)}")
    
    # Открывает окно для добавления нового расхода
    def open_add_expense_window(self): 
        self.expense_window = tk.Toplevel(self)
        self.expense_window.title("Добавить расход")
        self.expense_window.geometry("400x250")
        
        # Поля формы
        tk.Label(self.expense_window, text="Сумма:").grid(row=0, column=0, padx=10, pady=5, sticky='e')
        self.amount_entry = ttk.Entry(self.expense_window)
        self.amount_entry.grid(row=0, column=1, padx=10, pady=5)
        
        tk.Label(self.expense_window, text="Дата (ДД.ММ.ГГГГ):").grid(row=1, column=0, padx=10, pady=5, sticky='e')
        self.date_entry = ttk.Entry(self.expense_window)
        self.date_entry.grid(row=1, column=1, padx=10, pady=5)
        self.date_entry.insert(0, datetime.now().strftime("%d.%m.%Y"))
        
        tk.Label(self.expense_window, text="Описание:").grid(row=2, column=0, padx=10, pady=5, sticky='ne')
        self.desc_text = tk.Text(self.expense_window, width=30, height=5)
        self.desc_text.grid(row=2, column=1, padx=10, pady=5)
        
        ttk.Button(self.expense_window, text="Добавить", command=self.add_expense).grid(row=3, column=1, pady=10)
    
    # Добавляет новый расход в базу данных
    def add_expense(self):
        amount = self.amount_entry.get()
        date = self.date_entry.get()
        description = self.desc_text.get("1.0", tk.END).strip()
        
        # Валидация данных
        if not all([amount, date]):
            messagebox.showerror("Ошибка", "Сумма и дата обязательны для заполнения")
            return
        
        try:
            # Проверяем корректность суммы
            amount = float(amount)
            if amount <= 0:
                raise ValueError("Сумма должна быть положительной")
            
            # Проверяем корректность даты
            datetime.strptime(date, "%d.%m.%Y")
            
            # Добавляем запись в базу данных
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            cursor.execute(
                "INSERT INTO Finances (type, amount, date, description) VALUES (?, ?, ?, ?)",
                ("expense", amount, date, description)
            )
            conn.commit()
            conn.close()
            
            messagebox.showinfo("Успех", "Расход успешно добавлен")
            self.expense_window.destroy()
            self.generate_finance_report()
            
        except ValueError as e:
            messagebox.showerror("Ошибка", f"Некорректные данные: {str(e)}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось добавить расход: {str(e)}")
    
    # Генерирует финансовый отчет
    def generate_finance_report(self):
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        
        # Получаем доходы (из платежей)
        cursor.execute("SELECT SUM(amount) FROM Finances WHERE type='income'")
        total_income = cursor.fetchone()[0] or 0
        
        # Получаем расходы
        cursor.execute("SELECT SUM(amount) FROM Finances WHERE type='expense'")
        total_expense = cursor.fetchone()[0] or 0
        
        # Получаем все финансовые операции
        cursor.execute("""
            SELECT id, type, amount, date, description 
            FROM Finances 
            ORDER BY date DESC
        """)
        finance_data = cursor.fetchall()
        
        conn.close()
        
        # Обновляем таблицу
        for row in self.finance_tree.get_children():
            self.finance_tree.delete(row)
        
        for row in finance_data:
            item_id, type_, amount, date, description = row
            type_text = "Расход" if type_ == "expense" else "Доход"
            self.finance_tree.insert(
                "", "end", 
                values=(item_id, type_text, f"{amount:.2f} руб.", date, description)
            )
        
        # Обновляем итоговые значения
        self.income_label.config(text=f"{total_income:.2f} руб.")
        self.expense_label.config(text=f"{total_expense:.2f} руб.")
        self.total_label.config(text=f"{total_income - total_expense:.2f} руб.")

    # Удаляет выбранную финансовую запись
    def delete_finance_record(self):
        if not self.finance_tree.selection():
            messagebox.showwarning("Предупреждение", "Выберите запись для удаления")
            return
            
        record_id = self.finance_tree.item(self.finance_tree.selection()[0])['values'][0]
        
        if messagebox.askyesno("Подтверждение", "Вы действительно хотите удалить эту запись?"):
            try:
                conn = sqlite3.connect(get_db_path())
                cursor = conn.cursor()
                cursor.execute("DELETE FROM Finances WHERE id=?", (record_id,))
                conn.commit()
                conn.close()
                
                self.generate_finance_report()
                messagebox.showinfo("Успех", "Запись успешно удалена")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось удалить запись: {str(e)}")

    # Открывает окно для редактирования финансовой записи
    def open_edit_finance_window(self):
        if not self.finance_tree.selection():
            messagebox.showwarning("Предупреждение", "Выберите запись для редактирования")
            return
            
        record = self.finance_tree.item(self.finance_tree.selection()[0])['values']
        record_id, record_type, amount, date, description = record
        
        self.edit_finance_window = tk.Toplevel(self)
        self.edit_finance_window.title("Редактировать запись")
        self.edit_finance_window.geometry("400x300")
        
        # Поля формы
        ttk.Label(self.edit_finance_window, text="Тип:").grid(row=0, column=0, padx=10, pady=5, sticky='e')
        self.edit_type = ttk.Combobox(self.edit_finance_window, values=["Доход", "Расход"], state='readonly')
        self.edit_type.grid(row=0, column=1, padx=10, pady=5)
        self.edit_type.set(record_type)
        
        ttk.Label(self.edit_finance_window, text="Сумма:").grid(row=1, column=0, padx=10, pady=5, sticky='e')
        self.edit_amount = ttk.Entry(self.edit_finance_window)
        self.edit_amount.grid(row=1, column=1, padx=10, pady=5)
        self.edit_amount.insert(0, amount.split()[0])  # Убираем "руб."
        
        ttk.Label(self.edit_finance_window, text="Дата (ДД.ММ.ГГГГ):").grid(row=2, column=0, padx=10, pady=5, sticky='e')
        self.edit_date = ttk.Entry(self.edit_finance_window)
        self.edit_date.grid(row=2, column=1, padx=10, pady=5)
        self.edit_date.insert(0, date)
        
        ttk.Label(self.edit_finance_window, text="Описание:").grid(row=3, column=0, padx=10, pady=5, sticky='ne')
        self.edit_desc = tk.Text(self.edit_finance_window, width=30, height=5)
        self.edit_desc.grid(row=3, column=1, padx=10, pady=5)
        self.edit_desc.insert("1.0", description)
        
        ttk.Button(self.edit_finance_window, text="Сохранить", 
                  command=lambda: self.save_finance_record(record_id)).grid(row=4, column=1, pady=10)

    # Сохраняет изменения финансовой записи
    def save_finance_record(self, record_id):
        record_type = "income" if self.edit_type.get() == "Доход" else "expense"
        amount = self.edit_amount.get()
        date = self.edit_date.get()
        description = self.edit_desc.get("1.0", tk.END).strip()
        
        # Валидация данных
        if not all([amount, date]):
            messagebox.showerror("Ошибка", "Сумма и дата обязательны для заполнения")
            return
            
        try:
            amount = float(amount)
            datetime.strptime(date, "%d.%m.%Y")  # Проверка формата даты
        except ValueError:
            messagebox.showerror("Ошибка", "Некорректные данные (сумма должна быть числом, дата в формате ДД.ММ.ГГГГ)")
            return
            
        try:
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            cursor.execute("""
                UPDATE Finances 
                SET type=?, amount=?, date=?, description=?
                WHERE id=?
            """, (record_type, amount, date, description, record_id))
            conn.commit()
            conn.close()
            
            self.edit_finance_window.destroy()
            self.generate_finance_report()
            messagebox.showinfo("Успех", "Запись успешно обновлена")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось обновить запись: {str(e)}")

    
    # Поиск клиентов в отчете
    def search_clients(self):
        search_field = self.clients_search_field.get()
        search_text = self.clients_search_entry.get().strip().lower()
        
        if not search_text:
            self.generate_top_clients_report()
            return
            
        # Получаем все данные
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT 
                c.id, 
                c.name, 
                COALESCE((SELECT SUM(p.amount) FROM Payments p 
                        JOIN Reservations r ON p.reservation_id = r.id 
                        WHERE r.client_id = c.id), 0) +
                COALESCE((SELECT SUM(f.amount) 
                        FROM Finances f 
                        WHERE f.type='income' 
                        AND f.description LIKE '%Ручное добавление клиента ID:' || c.id || '%'), 0) as total_spent
            FROM Clients c
            GROUP BY c.id
            HAVING total_spent > 0
        """)
        
        all_clients = cursor.fetchall()
        conn.close()
        
        # Фильтруем результаты
        filtered_clients = []
        for row in all_clients:
            client_id, client_name, total_spent = row
            
            if search_field == "Клиент" and search_text in client_name.lower():
                filtered_clients.append(row)
            elif search_field == "Сумма":
                try:
                    search_amount = float(search_text)
                    if abs(total_spent - search_amount) < 0.01:  # Сравнение с учетом округления
                        filtered_clients.append(row)
                except ValueError:
                    pass
        
        # Обновляем таблицу
        for row in self.clients_tree.get_children():
            self.clients_tree.delete(row)
        
        for row in filtered_clients:
            client_id, client_name, total_spent = row
            self.clients_tree.insert("", "end", values=(client_id, client_name, f"{total_spent:.2f} руб."))
    
    # Сброс поиска клиентов
    def reset_clients_search(self):
        self.clients_search_entry.delete(0, tk.END)
        self.generate_top_clients_report()
    
    # Показывает детализацию платежей выбранного клиента
    def show_client_details(self, event):
        if not self.clients_tree.selection():
            return
            
        client_name = self.clients_tree.item(self.clients_tree.selection()[0])['values'][0]
        
        # Получаем ID клиента из имени
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM Clients WHERE name=?", (client_name,))
        client_id = cursor.fetchone()[0]
        
        # Получаем все платежи клиента
        cursor.execute("""
            SELECT p.id, r.room_id, rm.room_number, p.amount, p.payment_date
            FROM Payments p
            JOIN Reservations r ON p.reservation_id = r.id
            JOIN Rooms rm ON r.room_id = rm.id
            WHERE r.client_id=?
            ORDER BY p.payment_date DESC
        """, (client_id,))
        payments = cursor.fetchall()
        conn.close()
        
        # Создаем окно с детализацией
        details_window = tk.Toplevel(self)
        details_window.title(f"Платежи клиента {client_name}")
        details_window.geometry("600x400")
        
        # Таблица с платежами
        tree = ttk.Treeview(details_window, columns=("ID", "Номер", "Сумма", "Дата"), show="headings")
        for col in ("ID", "Номер", "Сумма", "Дата"):
            tree.heading(col, text=col)
            tree.column(col, width=100, anchor='center')
        
        for payment in payments:
            tree.insert("", "end", values=(payment[0], payment[2], f"{payment[3]:.2f} руб.", payment[4]))
        
        tree.pack(expand=True, fill='both', padx=5, pady=5)
        
        # Кнопка удаления платежа
        ttk.Button(details_window, text="Удалить платеж", 
                  command=lambda: self.delete_payment(tree, details_window)).pack(pady=5)

    # Показывает детализацию доходов по номеру
    def show_room_details(self, event):
        if not self.rooms_tree.selection():
            return
            
        room_number = self.rooms_tree.item(self.rooms_tree.selection()[0])['values'][0]
        
        # Получаем все платежи по номеру
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute("""
            SELECT p.id, c.name, p.amount, p.payment_date
            FROM Payments p
            JOIN Reservations r ON p.reservation_id = r.id
            JOIN Rooms rm ON r.room_id = rm.id
            JOIN Clients c ON r.client_id = c.id
            WHERE rm.room_number=?
            ORDER BY p.payment_date DESC
        """, (room_number,))
        payments = cursor.fetchall()
        conn.close()
        
        # Создаем окно с детализацией
        details_window = tk.Toplevel(self)
        details_window.title(f"Доходы номера {room_number}")
        details_window.geometry("600x400")
        
        # Таблица с платежами
        tree = ttk.Treeview(details_window, columns=("ID", "Клиент", "Сумма", "Дата"), show="headings")
        for col in ("ID", "Клиент", "Сумма", "Дата"):
            tree.heading(col, text=col)
            tree.column(col, width=100, anchor='center')
        
        for payment in payments:
            tree.insert("", "end", values=(payment[0], payment[1], f"{payment[2]:.2f} руб.", payment[3]))
        
        tree.pack(expand=True, fill='both', padx=5, pady=5)
        
        # Кнопка удаления платежа
        ttk.Button(details_window, text="Удалить платеж", 
                  command=lambda: self.delete_payment(tree, details_window)).pack(pady=5)

    # Удаляет выбранный платеж
    def delete_payment(self, tree, window):
        if not tree.selection():
            messagebox.showwarning("Предупреждение", "Выберите платеж для удаления")
            return
            
        payment_id = tree.item(tree.selection()[0])['values'][0]
        
        if messagebox.askyesno("Подтверждение", "Вы действительно хотите удалить этот платеж?"):
            try:
                conn = sqlite3.connect(get_db_path())
                cursor = conn.cursor()
                
                # Удаляем платеж
                cursor.execute("DELETE FROM Payments WHERE id=?", (payment_id,))
                
                # Удаляем соответствующую запись о доходе
                cursor.execute("""
                    DELETE FROM Finances 
                    WHERE type='income' AND description LIKE ?
                """, (f"%Оплата бронирования №{payment_id}%",))
                
                conn.commit()
                conn.close()
                
                # Обновляем данные
                window.destroy()
                self.generate_finance_report()
                self.generate_top_clients_report()
                self.generate_top_rooms_report()
                
                messagebox.showinfo("Успех", "Платеж успешно удален")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось удалить платеж: {str(e)}")

    # Генерирует отчет о лучших клиентах (включая ручные добавления)
    def generate_top_clients_report(self):
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        
        # Основной запрос: клиенты с бронированиями и ручными добавками
        cursor.execute("""
            SELECT 
                c.id, 
                c.name, 
                COALESCE((SELECT SUM(p.amount) FROM Payments p 
                        JOIN Reservations r ON p.reservation_id = r.id 
                        WHERE r.client_id = c.id), 0) +
                COALESCE((SELECT SUM(f.amount) 
                        FROM Finances f 
                        WHERE f.type='income' 
                        AND f.description LIKE '%Ручное добавление клиента ID:' || c.id || '%'), 0) as total_spent
            FROM Clients c
            GROUP BY c.id
            HAVING total_spent > 0
            ORDER BY total_spent DESC
            LIMIT 10
        """)
        
        # Очищаем таблицу
        for row in self.clients_tree.get_children():
            self.clients_tree.delete(row)
        
        # Добавляем клиентов из основного запроса
        for row in cursor.fetchall():
            client_id, client_name, total_spent = row
            self.clients_tree.insert("", "end", values=(client_id, client_name, f"{total_spent:.2f} руб."))
        
        # Отдельный запрос для клиентов без бронирований
        cursor.execute("""
            SELECT 
                c.id as client_id,
                c.name as client_name,
                SUM(f.amount) as total_spent
            FROM Finances f
            JOIN Clients c ON f.description LIKE '%Ручное добавление клиента ID:' || c.id || '%'
            WHERE f.type='income'
            AND NOT EXISTS (
                SELECT 1 FROM Reservations r WHERE r.client_id = c.id
            )
            GROUP BY c.id
            HAVING total_spent > 0
            ORDER BY total_spent DESC
        """)
        
        # Добавляем клиентов без бронирований, но с ручными платежами
        for row in cursor.fetchall():
            client_id, client_name, total_spent = row
            # проверка, нет ли уже этого клиента в таблице
            if not any(self.clients_tree.item(item)['values'][0] == client_id for item in self.clients_tree.get_children()):
                self.clients_tree.insert("", "end", values=(client_id, client_name, f"{total_spent:.2f} руб."))
        
        conn.close()

    # Поиск номеров в отчете
    def search_rooms_report(self):
        search_field = self.rooms_search_field.get()
        search_text = self.rooms_search_entry.get().strip().lower()
        
        if not search_text:
            self.generate_top_rooms_report()
            return
            
        # Получаем все данные
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT r.id, r.room_number, SUM(p.amount) as total_income
            FROM Rooms r
            JOIN Reservations res ON r.id = res.room_id
            JOIN Payments p ON res.id = p.reservation_id
            WHERE res.client_id != 1  -- Исключаем тестовые бронирования
            GROUP BY r.id
            ORDER BY total_income DESC
        """)
        
        all_rooms = cursor.fetchall()
        conn.close()
        
        # Фильтруем результаты
        filtered_rooms = []
        for row in all_rooms:
            room_id, room_number, total_income = row
            
            if search_field == "Номер" and search_text in room_number.lower():
                filtered_rooms.append(row)
            elif search_field == "Доход":
                try:
                    search_amount = float(search_text)
                    if abs(total_income - search_amount) < 0.01:  # Сравнение с учетом округления
                        filtered_rooms.append(row)
                except ValueError:
                    pass
        
        # Обновляем таблицу
        for row in self.rooms_tree.get_children():
            self.rooms_tree.delete(row)
        
        for row in filtered_rooms:
            room_id, room_number, total_income = row
            self.rooms_tree.insert("", "end", values=(room_id, room_number, f"{total_income:.2f} руб."))
    
    # Сброс поиска номеров
    def reset_rooms_search(self):
        self.rooms_search_entry.delete(0, tk.END)
        self.generate_top_rooms_report()
    
    # Генерирует отчет о лучших номерах (исключая тестовые бронирования)
    def generate_top_rooms_report(self):
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT r.id, r.room_number, SUM(p.amount) as total_income
            FROM Rooms r
            JOIN Reservations res ON r.id = res.room_id
            JOIN Payments p ON res.id = p.reservation_id
            WHERE res.client_id != 1  -- Исключаем тестовые бронирования
            GROUP BY r.id
            ORDER BY total_income DESC
            LIMIT 10
        """)
        
        # Очищаем таблицу
        for row in self.rooms_tree.get_children():
            self.rooms_tree.delete(row)
        
        # Заполняем данными
        for row in cursor.fetchall():
            room_id, room_number, total_income = row
            self.rooms_tree.insert("", "end", values=(room_id, room_number, f"{total_income:.2f} руб."))
        
        conn.close()
    
    # Окно для добавления клиента в отчет
    def open_add_client_window(self):
        self.add_client_window = tk.Toplevel(self)
        self.add_client_window.title("Добавить клиента")
        self.add_client_window.geometry("400x200")
        
        # Получаем список клиентов из базы
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute("SELECT id, name FROM Clients ORDER BY name")
        clients = cursor.fetchall()
        conn.close()
        
        ttk.Label(self.add_client_window, text="Клиент:").pack(pady=5)
        self.client_combobox = ttk.Combobox(self.add_client_window, values=[f"{c[1]} (ID: {c[0]})" for c in clients])
        self.client_combobox.pack(pady=5)
        
        ttk.Label(self.add_client_window, text="Сумма:").pack(pady=5)
        self.client_amount_entry = ttk.Entry(self.add_client_window)
        self.client_amount_entry.pack(pady=5)
        
        ttk.Button(self.add_client_window, text="Добавить", command=self.add_client_to_report).pack(pady=10)
    
    # Добавляет клиента в отчет и сразу обновляет список
    def add_client_to_report(self):
        client_str = self.client_combobox.get()
        amount = self.client_amount_entry.get()
        
        if not all([client_str, amount]):
            messagebox.showerror("Ошибка", "Заполните все поля")
            return
            
        try:
            client_id = int(client_str.split("(ID: ")[1][:-1])
            amount = float(amount)
            
            if amount <= 0:
                messagebox.showerror("Ошибка", "Сумма должна быть положительной")
                return
                
            # Добавляем платеж в базу
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            
            # 1. Добавляем запись в Finances
            cursor.execute("""
                INSERT INTO Finances (type, amount, date, description)
                VALUES ('income', ?, date('now'), ?)
            """, (amount, f"Ручное добавление клиента ID: {client_id}"))
            
            # 2. Проверяем, что клиент существует
            cursor.execute("SELECT name FROM Clients WHERE id=?", (client_id,))
            if not cursor.fetchone():
                messagebox.showerror("Ошибка", "Клиент с таким ID не найден")
                conn.rollback()
                conn.close()
                return
                
            conn.commit()
            conn.close()
            
            # 3. Обновляем список клиентов
            self.generate_top_clients_report()
            
            # 4. Закрываем окно и показываем сообщение
            self.add_client_window.destroy()
            messagebox.showinfo("Успех", "Клиент успешно добавлен в отчет")
            
        except ValueError:
            messagebox.showerror("Ошибка", "Введите корректные данные (числа)")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось добавить клиента: {str(e)}")
            if 'conn' in locals():
                conn.rollback()
                conn.close()

    # Удаляет клиента из отчета (удаляет связанные платежи)
    def delete_client(self):
        if not self.clients_tree.selection():
            messagebox.showwarning("Предупреждение", "Выберите клиента для удаления")
            return
            
        client_id = self.clients_tree.item(self.clients_tree.selection()[0])['values'][0]
        client_name = self.clients_tree.item(self.clients_tree.selection()[0])['values'][1]
        
        if messagebox.askyesno("Подтверждение", f"Удалить все платежи клиента {client_name} из отчета?"):
            try:
                conn = sqlite3.connect(get_db_path())
                cursor = conn.cursor()
                
                cursor.execute("""
                    DELETE FROM Finances 
                    WHERE description LIKE ? AND type='income'
                """, (f"%Ручное добавление клиента ID: {client_id}%",))
                
                conn.commit()
                conn.close()
                
                self.generate_top_clients_report()
                messagebox.showinfo("Успех", "Платежи клиента удалены из отчета")
                
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось удалить платежи: {str(e)}")
    
    # Окно для добавления номера в отчет
    def open_add_room_window(self):
        self.add_room_window = tk.Toplevel(self)
        self.add_room_window.title("Добавить номер")
        self.add_room_window.geometry("400x200")
        
        # Получаем список номеров из базы
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute("SELECT id, room_number FROM Rooms ORDER BY room_number")
        rooms = cursor.fetchall()
        conn.close()
        
        ttk.Label(self.add_room_window, text="Номер:").pack(pady=5)
        self.room_combobox = ttk.Combobox(self.add_room_window, values=[f"{r[1]} (ID: {r[0]})" for r in rooms])
        self.room_combobox.pack(pady=5)
        
        ttk.Label(self.add_room_window, text="Сумма:").pack(pady=5)
        self.room_amount_entry = ttk.Entry(self.add_room_window)
        self.room_amount_entry.pack(pady=5)
        
        ttk.Button(self.add_room_window, text="Добавить", command=self.add_room_to_report).pack(pady=10)
    
    # Добавляет номер в отчет
    def add_room_to_report(self):
        room_str = self.room_combobox.get()
        amount = self.room_amount_entry.get()
        
        if not all([room_str, amount]):
            messagebox.showerror("Ошибка", "Заполните все поля")
            return
            
        try:
            room_id = int(room_str.split("(ID: ")[1][:-1])
            amount = float(amount)
            
            # Добавляем платеж в базу
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            
            cursor.execute("""
                INSERT INTO Finances (type, amount, date, description)
                VALUES ('income', ?, date('now'), ?)
            """, (amount, f"Ручное добавление номера ID: {room_id} в отчет"))
            
            conn.commit()
            conn.close()
            
            self.add_room_window.destroy()
            self.generate_top_rooms_report()
            messagebox.showinfo("Успех", "Номер успешно добавлен в отчет")
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось добавить номер: {str(e)}")
    
    # Удаляет номер из отчета (удаляет связанные платежи)
    def delete_room(self):
        if not self.rooms_tree.selection():
            messagebox.showwarning("Предупреждение", "Выберите номер для удаления")
            return
            
        room_id = self.rooms_tree.item(self.rooms_tree.selection()[0])['values'][0]
        room_number = self.rooms_tree.item(self.rooms_tree.selection()[0])['values'][1]
        
        if messagebox.askyesno("Подтверждение", f"Удалить все платежи номера {room_number} из отчета?"):
            try:
                conn = sqlite3.connect(get_db_path())
                cursor = conn.cursor()
                
                cursor.execute("""
                    DELETE FROM Finances 
                    WHERE description LIKE ? AND type='income'
                """, (f"%Ручное добавление номера ID: {room_id}%",))

                conn.commit()
                conn.close()
                
                self.generate_top_rooms_report()
                messagebox.showinfo("Успех", "Платежи номера удалены из отчета")
                
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось удалить платежи: {str(e)}")

    def export_current_tab_to_excel(self):
        current_tab = self.notebook.tab(self.notebook.select(), "text")
        self.export_to_excel(current_tab)

    # Экспортирует текущую активную вкладку в Word
    def export_current_tab_to_word(self):
        current_tab = self.notebook.tab(self.notebook.select(), "text")
        self.export_to_word(current_tab)

    # Экспортирует отчет в Excel
    def export_to_excel(self, report_type):
        if report_type == "Финансы":
            data = [self.finance_tree.item(item)['values'] for item in self.finance_tree.get_children()]
            headers = ["ID", "Тип", "Сумма", "Дата", "Описание"]
        elif report_type == "Лучшие клиенты":
            data = [self.clients_tree.item(item)['values'] for item in self.clients_tree.get_children()]
            headers = ["Клиент", "Потрачено"]
        else:  # Лучшие номера
            data = [self.rooms_tree.item(item)['values'] for item in self.rooms_tree.get_children()]
            headers = ["Номер", "Доход"]
        
        if not data:
            messagebox.showwarning("Предупреждение", "Нет данных для экспорта")
            return
        
        # Создаем диалог выбора файла
        file_path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel файлы", "*.xlsx"), ("Все файлы", "*.*")],
            title="Сохранить отчет как"
        )
        
        if not file_path:  # Пользователь отменил сохранение
            return
        
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = report_type
            
            # Добавляем заголовки
            ws.append(headers)
            
            # Добавляем данные
            for row in data:
                ws.append(row)
            
            # итоговые суммы для финансового отчета
            if report_type == "Финансы":
                # Рассчитываем суммы
                total_income = sum(float(row[2].split()[0]) for row in data if row[1] == "Доход")
                total_expense = sum(float(row[2].split()[0]) for row in data if row[1] == "Расход")
                profit = total_income - total_expense
                
                # Добавляем итоги
                ws.append([])  # Пустая строка
                ws.append(["Общий доход:", f"{total_income:.2f} руб."])
                ws.append(["Общий расход:", f"{total_expense:.2f} руб."])
                ws.append(["Прибыль:", f"{profit:.2f} руб."])
            
            # Настраиваем стиль
            header_font = openpyxl.styles.Font(bold=True)
            for cell in ws[1]:
                cell.font = header_font
            
            # Автоматическая ширина столбцов
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = (max_length + 2)
                ws.column_dimensions[column_letter].width = adjusted_width
            
            # Добавляем текущую дату и время
            now = datetime.now().strftime("%d.%m.%Y %H:%M")
            ws.cell(row=ws.max_row+2, column=1, value=f"Отчет сгенерирован {now}")
            
            wb.save(file_path)
            messagebox.showinfo("Успех", f"Отчет успешно сохранен в файл:\n{file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить файл: {str(e)}")

    # Экспортирует отчет в Word
    def export_to_word(self, report_type):
        if report_type == "Финансы":
            data = [self.finance_tree.item(item)['values'] for item in self.finance_tree.get_children()]
            headers = ["ID", "Тип", "Сумма", "Дата", "Описание"]
        elif report_type == "Лучшие клиенты":
            data = [self.clients_tree.item(item)['values'] for item in self.clients_tree.get_children()]
            headers = ["Клиент", "Потрачено"]
        else:  # Лучшие номера
            data = [self.rooms_tree.item(item)['values'] for item in self.rooms_tree.get_children()]
            headers = ["Номер", "Доход"]
        
        if not data:
            messagebox.showwarning("Предупреждение", "Нет данных для экспорта")
            return
        
        # Создаем диалог выбора файла
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word файлы", "*.docx"), ("Все файлы", "*.*")],
            title="Сохранить отчет как"
        )
        
        if not file_path: 
            return
        
        try:
            from docx import Document
            from docx.shared import Pt
            
            doc = Document()
            
            # Добавляем заголовок
            doc.add_heading(f'Отчет: {report_type}', level=1)
            
            # Добавляем таблицу
            table = doc.add_table(rows=1, cols=len(headers))
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
            
            # Добавляем данные
            for row in data:
                row_cells = table.add_row().cells
                for i, value in enumerate(row[1:] if report_type != "Финансы" else row):  # Для финансов показываем все колонки
                    row_cells[i].text = str(value)
            
            # Добавляем итоговые суммы для финансового отчета
            if report_type == "Финансы":
                # Рассчитываем суммы
                total_income = sum(float(row[2].split()[0]) for row in data if row[1] == "Доход")
                total_expense = sum(float(row[2].split()[0]) for row in data if row[1] == "Расход")
                profit = total_income - total_expense
                
                # Добавляем итоги
                doc.add_paragraph()
                doc.add_paragraph(f"Общий доход: {total_income:.2f} руб.")
                doc.add_paragraph(f"Общий расход: {total_expense:.2f} руб.")
                doc.add_paragraph(f"Прибыль: {profit:.2f} руб.")
            
            # Добавляем дату и время
            now = datetime.now().strftime("%d.%m.%Y %H:%M")
            doc.add_paragraph(f"Отчет сгенерирован {now}")
            
            doc.save(file_path)
            messagebox.showinfo("Успех", f"Отчет успешно сохранен в файл:\n{file_path}")
        except ImportError:
            messagebox.showerror("Ошибка", "Для экспорта в Word требуется установить библиотеку python-docx")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить файл: {str(e)}")

# --- Класс, описывающий создание фрейма ОПЛАТА ---
class Payment(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        testLbl = ttk.Label(self)
        testLbl.pack(expand=True, fill='both')
        self.place(relx=0.3, y=0, relwidth=0.7, relheight=1)
        self.create_widgets()
        self.fetch_clients()
        self.fetch_reservations()

    def create_widgets(self):
        # Основной контейнер
        main_frame = ttk.Frame(self, padding=10)
        main_frame.pack(fill='both', expand=True)
        
        # Выбор клиента
        ttk.Label(main_frame, text="Клиент:").grid(row=0, column=0, padx=5, pady=5, sticky='e')
        self.client_combobox = ttk.Combobox(main_frame, state='readonly')
        self.client_combobox.grid(row=0, column=1, padx=5, pady=5, sticky='ew')
        self.client_combobox.bind('<<ComboboxSelected>>', self.update_reservations)
        
        # Выбор бронирования
        ttk.Label(main_frame, text="Бронирование:").grid(row=1, column=0, padx=5, pady=5, sticky='e')
        self.reservation_combobox = ttk.Combobox(main_frame, state='readonly')
        self.reservation_combobox.grid(row=1, column=1, padx=5, pady=5, sticky='ew')
        self.reservation_combobox.bind('<<ComboboxSelected>>', self.show_reservation_details)
        
        # Информация о бронировании
        self.reservation_info = ttk.LabelFrame(main_frame, text="Информация о бронировании", padding=10)
        self.reservation_info.grid(row=2, column=0, columnspan=2, padx=5, pady=5, sticky='nsew')
        
        ttk.Label(self.reservation_info, text="Номер:").grid(row=0, column=0, sticky='e')
        self.room_label = ttk.Label(self.reservation_info, text="")
        self.room_label.grid(row=0, column=1, sticky='w')
        
        ttk.Label(self.reservation_info, text="Даты:").grid(row=1, column=0, sticky='e')
        self.dates_label = ttk.Label(self.reservation_info, text="")
        self.dates_label.grid(row=1, column=1, sticky='w')
        
        ttk.Label(self.reservation_info, text="Сумма к оплате:").grid(row=2, column=0, sticky='e')
        self.amount_label = ttk.Label(self.reservation_info, text="", font=('Arial', 10, 'bold'))
        self.amount_label.grid(row=2, column=1, sticky='w')
        
        # Поле для ввода суммы оплаты
        ttk.Label(main_frame, text="Сумма оплаты:").grid(row=3, column=0, padx=5, pady=5, sticky='e')
        self.payment_entry = ttk.Entry(main_frame)
        self.payment_entry.grid(row=3, column=1, padx=5, pady=5, sticky='ew')
        
        # Кнопки
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=4, column=0, columnspan=2, pady=10)
        
        ttk.Button(button_frame, text="Сформировать чек", command=self.generate_receipt).pack(side='left', padx=5)
        ttk.Button(button_frame, text="Обновить", command=self.refresh_data).pack(side='right', padx=5)
        
        # Настройка веса строк и столбцов
        main_frame.columnconfigure(1, weight=1)
        self.reservation_info.columnconfigure(1, weight=1)

    def fetch_clients(self):
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute("SELECT id, name FROM Clients ORDER BY name")
        self.clients = cursor.fetchall()
        self.client_combobox['values'] = [f"{c[1]} (ID: {c[0]})" for c in self.clients]
        conn.close()

    def fetch_reservations(self):
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute("""
            SELECT r.id, r.client_id, r.room_id, rm.room_number, 
                   r.checkin_date, r.checkout_date, r.total_price
            FROM Reservations r
            JOIN Rooms rm ON r.room_id = rm.id
            WHERE r.checkout_date >= date('now')
            ORDER BY r.checkin_date
        """)
        self.reservations = cursor.fetchall()
        conn.close()

    def update_reservations(self, event):
        if not self.client_combobox.get():
            return
            
        # Получаем ID клиента
        client_str = self.client_combobox.get()
        client_id = int(client_str.split("(ID: ")[1][:-1])
        
        # Фильтруем бронирования для этого клиента
        client_reservations = [r for r in self.reservations if r[1] == client_id]
        
        # Обновляем комбобокс бронирований
        self.reservation_combobox['values'] = [
            f"№{r[3]} с {r[4]} по {r[5]} (ID: {r[0]})" 
            for r in client_reservations
        ]
        
        # Очищаем информацию о бронировании
        self.room_label.config(text="")
        self.dates_label.config(text="")
        self.amount_label.config(text="")

    def show_reservation_details(self, event):
        if not self.reservation_combobox.get():
            return
            
        # Получаем ID бронирования
        res_str = self.reservation_combobox.get()
        res_id = int(res_str.split("(ID: ")[1][:-1])
        
        # Находим бронирование
        reservation = next(r for r in self.reservations if r[0] == res_id)
        
        # Обновляем информацию
        self.room_label.config(text=f"{reservation[3]} (ID: {reservation[2]})")
        self.dates_label.config(text=f"{reservation[4]} - {reservation[5]}")
        self.amount_label.config(text=f"{reservation[6]:.2f} руб.")
        self.payment_entry.delete(0, tk.END)
        self.payment_entry.insert(0, str(reservation[6]))
        
    def pay_reservation(self):
        if not self.selected_reservation:
            messagebox.showwarning("Предупреждение", "Выберите бронирование для оплаты")
            return
            
        reservation_id = self.reservations_tree.item(self.selected_reservation)['values'][0]
        
        # Проверяем, не оплачена ли уже бронь
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        cursor.execute("SELECT 1 FROM Payments WHERE reservation_id=?", (reservation_id,))
        if cursor.fetchone():
            messagebox.showwarning("Предупреждение", "Это бронирование уже оплачено")
            conn.close()
            return

    def generate_receipt(self):
        if not self.reservation_combobox.get():
            messagebox.showerror("Ошибка", "Выберите бронирование")
            return
            
        try:
            amount = float(self.payment_entry.get())
            if amount <= 0:
                raise ValueError("Сумма должна быть положительной")
        except ValueError:
            messagebox.showerror("Ошибка", "Введите корректную сумму")
            return
            
        # Получаем ID бронирования
        res_str = self.reservation_combobox.get()
        res_id = int(res_str.split("(ID: ")[1][:-1])
        
        # Получаем данные бронирования
        reservation = next(r for r in self.reservations if r[0] == res_id)
        
        # Создаем чек
        receipt_text = f"""
        ЧЕК ОБ ОПЛАТЕ
        ----------------------------
        Клиент: {self.client_combobox.get()}
        Номер: {reservation[3]}
        Период: {reservation[4]} - {reservation[5]}
        Сумма к оплате: {reservation[6]:.2f} руб.
        Оплачено: {amount:.2f} руб.
        Дата оплаты: {datetime.now().strftime("%d.%m.%Y %H:%M")}
        ----------------------------
        Спасибо за выбор нашего отеля!
        """
        
        messagebox.showinfo("Чек об оплате", receipt_text)
        
        # Записываем платеж в базу данных
        try:
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            
            # Добавляем платеж
            cursor.execute("""
                INSERT INTO Payments (reservation_id, amount, payment_date)
                VALUES (?, ?, ?)
            """, (res_id, amount, datetime.now().strftime("%Y-%m-%d")))
            
            # Добавляем доход в таблицу Finances
            cursor.execute("""
                INSERT INTO Finances (type, amount, date, description)
                VALUES (?, ?, ?, ?)
            """, ("income", amount, datetime.now().strftime("%Y-%m-%d"), 
                f"Оплата бронирования №{res_id}"))
            
            conn.commit()
            conn.close()
            
            messagebox.showinfo("Успех", "Платеж успешно зарегистрирован")
            
            # Обновляем отчеты
            self.master.report.generate_finance_report()
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить платеж: {str(e)}")
    
    def refresh_data(self):
        self.fetch_clients()
        self.fetch_reservations()
        self.client_combobox.set('')
        self.reservation_combobox.set('')
        self.room_label.config(text="")
        self.dates_label.config(text="")
        self.amount_label.config(text="")
        self.payment_entry.delete(0, tk.END)

# --- Класс, описывающий создание фрейма АДМИНИСТРИРОВАНИЕ ---
class AdminPanel(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        testLbl = ttk.Label(self, background='#35A7FF')
        testLbl.pack(expand=True, fill='both')
        self.place(relx=0.3, y=0, relwidth=0.7, relheight=1)
        self.create_widgets()
        self.backup_running = False
        self.start_backup_scheduler()
        
    def create_widgets(self):
        # Основной контейнер
        main_frame = ttk.Frame(self)
        main_frame.place(relx=0.05, rely=0.05, relwidth=0.9, relheight=0.9)
        
        # Вкладки
        self.notebook = ttk.Notebook(main_frame)
        self.notebook.pack(expand=True, fill='both')
        
        # Вкладка управления пользователями
        self.users_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.users_tab, text="Управление пользователями")
        self.create_users_tab()
        
        # Вкладка резервного копирования
        self.backup_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.backup_tab, text="Резервное копирование")
        self.create_backup_tab()
    
    def create_users_tab(self):
        # Таблица пользователей
        # Таблица пользователей
        self.users_tree = ttk.Treeview(
            self.users_tab, 
            columns=("ID", "Логин"), 
            show="headings"
        )
        self.users_tree.heading("ID", text="ID")
        self.users_tree.heading("Логин", text="Логин")
        
        self.users_tree.column("ID", width=50, anchor='center')
        self.users_tree.column("Логин", width=250, anchor='w')
        
        self.users_tree.pack(expand=True, fill='both', padx=5, pady=5)
        
        # Кнопки управления
        button_frame = ttk.Frame(self.users_tab)
        button_frame.pack(fill='x', padx=5, pady=5)
        
        ttk.Button(button_frame, text="Обновить", command=self.load_users).pack(side='left', padx=5)
        ttk.Button(button_frame, text="Добавить пользователя", command=self.open_add_user_window).pack(side='left', padx=5)
        ttk.Button(button_frame, text="Изменить пароль", command=self.open_change_password_window).pack(side='left', padx=5)
        ttk.Button(button_frame, text="Удалить", command=self.delete_user).pack(side='left', padx=5)
        
        self.load_users()
    
    # Загружает список пользователей из базы данных
    def load_users(self):
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        
        cursor.execute("SELECT id, name FROM Users ORDER BY id")
        users = cursor.fetchall()
        conn.close()
        
        # Очищаем таблицу
        for row in self.users_tree.get_children():
            self.users_tree.delete(row)
        
        # Заполняем данными
        for user in users:
            self.users_tree.insert("", "end", values=user)
    
    # Открывает окно для добавления нового пользователя
    def open_add_user_window(self):
        self.add_user_window = tk.Toplevel(self)
        self.add_user_window.title("Добавить пользователя")
        self.add_user_window.geometry("300x200")
        
        ttk.Label(self.add_user_window, text="Логин:").pack(pady=5)
        self.new_username = ttk.Entry(self.add_user_window)
        self.new_username.pack(pady=5)
        
        ttk.Label(self.add_user_window, text="Пароль:").pack(pady=5)
        self.new_password = ttk.Entry(self.add_user_window, show="*")
        self.new_password.pack(pady=5)
        
        ttk.Button(self.add_user_window, text="Добавить", command=self.add_user).pack(pady=10)
    
    # Добавляет нового пользователя в базу данных
    def add_user(self):
        username = self.new_username.get().strip()
        password = self.new_password.get().strip()
        
        if not username or not password:
            messagebox.showerror("Ошибка", "Логин и пароль не могут быть пустыми")
            return
            
        try:
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            
            # Проверяем, существует ли уже пользователь с таким логином
            cursor.execute("SELECT id FROM Users WHERE name=?", (username,))
            if cursor.fetchone():
                messagebox.showerror("Ошибка", "Пользователь с таким логином уже существует")
                return
            
            # Добавляем нового пользователя
            cursor.execute("INSERT INTO Users (name, password) VALUES (?, ?)", (username, password))
            conn.commit()
            
            messagebox.showinfo("Успех", "Пользователь успешно добавлен")
            self.add_user_window.destroy()
            self.load_users()
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось добавить пользователя: {str(e)}")
        finally:
            if 'conn' in locals():
                conn.close()
    
    # Открывает окно для изменения пароля пользователя
    def open_change_password_window(self):
        if not self.users_tree.selection():
            messagebox.showwarning("Предупреждение", "Выберите пользователя")
            return
            
        user_id = self.users_tree.item(self.users_tree.selection()[0])['values'][0]
        username = self.users_tree.item(self.users_tree.selection()[0])['values'][1]
        
        # Защищаем учетную запись root от изменений
        if username.lower() == "root":
            messagebox.showerror("Ошибка", "Нельзя изменять пароль суперпользователя root")
            return
        
        self.change_pass_window = tk.Toplevel(self)
        self.change_pass_window.title(f"Изменить пароль для {username}")
        self.change_pass_window.geometry("300x200")
        
        ttk.Label(self.change_pass_window, text="Новый пароль:").pack(pady=5)
        self.new_pass_entry = ttk.Entry(self.change_pass_window, show="*")
        self.new_pass_entry.pack(pady=5)
        
        ttk.Label(self.change_pass_window, text="Подтвердите пароль:").pack(pady=5)
        self.confirm_pass_entry = ttk.Entry(self.change_pass_window, show="*")
        self.confirm_pass_entry.pack(pady=5)
        
        self.changing_user_id = user_id
        
        ttk.Button(self.change_pass_window, text="Сохранить", command=self.change_password).pack(pady=10)
    
    # Изменяет пароль выбранного пользователя
    def change_password(self):
        new_pass = self.new_pass_entry.get()
        confirm_pass = self.confirm_pass_entry.get()
        
        if not new_pass or not confirm_pass:
            messagebox.showerror("Ошибка", "Пароль не может быть пустым")
            return
            
        if new_pass != confirm_pass:
            messagebox.showerror("Ошибка", "Пароли не совпадают")
            return
            
        try:
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            
            cursor.execute("UPDATE Users SET password=? WHERE id=?", (new_pass, self.changing_user_id))
            conn.commit()
            
            messagebox.showinfo("Успех", "Пароль успешно изменен")
            self.change_pass_window.destroy()
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось изменить пароль: {str(e)}")
        finally:
            if 'conn' in locals():
                conn.close()
    
    # Удаляет выбранного пользователя
    def delete_user(self):
        if not self.users_tree.selection():
            messagebox.showwarning("Предупреждение", "Выберите пользователя")
            return
            
        user_id = self.users_tree.item(self.users_tree.selection()[0])['values'][0]
        username = self.users_tree.item(self.users_tree.selection()[0])['values'][1]
        
        # Защищаем учетную запись root от удаления
        if username.lower() == "root":
            messagebox.showerror("Ошибка", "Нельзя удалить суперпользователя root")
            return
            
        if not messagebox.askyesno("Подтверждение", f"Вы действительно хотите удалить пользователя {username}?"):
            return
            
        try:
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            
            cursor.execute("DELETE FROM Users WHERE id=?", (user_id,))
            conn.commit()
            
            messagebox.showinfo("Успех", "Пользователь успешно удален")
            self.load_users()
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось удалить пользователя: {str(e)}")
        finally:
            if 'conn' in locals():
                conn.close()
    
    def create_backup_tab(self):
        # Информация о последнем бэкапе
        self.last_backup_label = ttk.Label(self.backup_tab, text="Последнее резервное копирование: никогда")
        self.backup_status_label = ttk.Label(self.backup_tab, text="Статус: не активен")
        
        self.last_backup_label.pack(pady=10)
        self.backup_status_label.pack(pady=5)
        
        # Кнопка ручного бэкапа
        ttk.Button(
            self.backup_tab, 
            text="Создать резервную копию сейчас", 
            command=self.create_backup_manual
        ).pack(pady=10)
        
        # История бэкапов
        ttk.Label(self.backup_tab, text="История резервных копий:").pack(pady=5)
        
        self.backup_tree = ttk.Treeview(
            self.backup_tab, 
            columns=("Дата", "Путь", "Статус"), 
            show="headings"
        )
        self.backup_tree.heading("Дата", text="Дата")
        self.backup_tree.heading("Путь", text="Путь")
        self.backup_tree.heading("Статус", text="Статус")
        
        self.backup_tree.column("Дата", width=150, anchor='center')
        self.backup_tree.column("Путь", width=250, anchor='w')
        self.backup_tree.column("Статус", width=100, anchor='center')
        
        self.backup_tree.pack(expand=True, fill='both', padx=5, pady=5)
        
        # Загружаем историю бэкапов
        self.load_backup_history()
    
    # Запускает планировщик автоматического резервного копирования
    def start_backup_scheduler(self):
        if not self.backup_running:
            self.backup_running = True
            self.schedule_backup()
    
    # Планирует следующее резервное копирование
    def schedule_backup(self):
        # Запускаем бэкап через час (3600000 мс)
        self.after(3600000, self.create_backup_auto)
        #self.after(24000, self.create_backup_auto) # -- test
        
    # Создает автоматическую резервную копию
    def create_backup_auto(self):
        try:
            self.backup_status_label.config(text="Статус: создание резервной копии...")
            self.update()
            
            backup_dir = "D:/backups"
            if not os.path.exists(backup_dir):
                os.makedirs(backup_dir)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_path = os.path.join(backup_dir, f"hotel_backup_{timestamp}.db")
            
            if self.backup_database('hotel.db', backup_path):
                self.last_backup_label.config(text=f"Последнее резервное копирование: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
                self.backup_status_label.config(text="Статус: успешно завершено")
                self.backup_tree.insert("", "end", values=(
                    datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
                    backup_path,
                    "Успешно"
                ))
            else:
                raise Exception("Не удалось создать резервную копию")
                
            self.schedule_backup()
            
        except Exception as e:
            self.backup_status_label.config(text=f"Статус: ошибка ({str(e)}), повтор через 2 минуты")
            self.after(120000, self.create_backup_auto)
    
    # Проверяет, является ли файл валидной SQLite базой 
    def is_valid_backup(self, file_path):
        try:
            conn = sqlite3.connect(file_path)
            conn.execute("PRAGMA quick_check;")
            conn.close()
            return True
        except:
            return False
    
    def create_backup_manual(self):
        try:
            # Обновляем статус в интерфейсе
            self.backup_status_label.config(text="Статус: создание резервной копии...")
            self.update()  # Принудительно обновляем интерфейс
            
            # Создаем папку для бэкапов, если ее нет
            backup_dir = "D:/backups"
            if not os.path.exists(backup_dir):
                os.makedirs(backup_dir)
            
            # Формируем имя файла с датой и временем
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_path = os.path.join(backup_dir, f"hotel_manual_backup_{timestamp}.db")
            
            # Подключаемся к исходной и целевой базам
            source_conn = sqlite3.connect(get_db_path())
            target_conn = sqlite3.connect(backup_path)
            
            # Выполняем резервное копирование
            source_conn.backup(target_conn)
            
            # Закрываем соединения
            source_conn.close()
            target_conn.close()
            
            # Обновляем интерфейс
            self.last_backup_label.config(
                text=f"Последнее резервное копирование: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}"
            )
            #self.backup_status_label.config(text="Статус: успешно завершено")
            
            is_valid = self.is_valid_backup(backup_path)
            status = "Успешно (ручной)" if is_valid else "Ошибка (ручной)"
            
            self.backup_tree.insert("", "end", values=(
                datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
                backup_path,
                status
            ))
            
            messagebox.showinfo(
                "Успех", 
                f"Резервная копия успешно создана:\n{backup_path}"
            )
            
            
            
        except Exception as e:
            # В случае ошибки сразу пишем "Ошибка"
            self.backup_tree.insert("", "end", values=(
                datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
                backup_path,
                "Ошибка (ручной)"
            ))
            
            # Если соединения были открыты - закрываем их
            if 'target_conn' in locals():
                target_conn.close()
            if 'source_conn' in locals():
                source_conn.close()
    
    def backup_database(self, source_db, target_db):
        try:
            
            source_conn = sqlite3.connect(source_db)
            target_conn = sqlite3.connect(target_db)
            source_conn.backup(target_conn)
            
            messagebox.showinfo("Успех", f"Резервная копия успешно создана:\n{target_db}")
            return True
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось создать резервную копию:\n{str(e)}")
            return False
            
        finally:
            # закрываем соединения
            if 'target_conn' in locals():
                target_conn.close()
            if 'source_conn' in locals():
                source_conn.close()
    
    # Загружает историю резервных копий
    def load_backup_history(self):
        backup_dir = "D:/backups"
        if not os.path.exists(backup_dir):
            return
            
        backup_files = glob.glob(os.path.join(backup_dir, "hotel_*.db"))
        backup_files.sort(key=os.path.getmtime, reverse=True)
        
        for file_path in backup_files[:20]:  # Первые 20 файлов
            mtime = os.path.getmtime(file_path)
            date_str = datetime.fromtimestamp(mtime).strftime("%d.%m.%Y %H:%M:%S")
            
            # Определение типа бэкапа
            is_manual = "manual" in file_path.lower()
            
            # Проверка целостности файла
            is_valid = self.is_valid_backup(file_path)
            
            # Формирование статуса
            if is_manual:
                status = "Успешно (ручной)" if is_valid else "Ошибка (ручной)"
            else:
                status = "Успешно (авто)" if is_valid else "Ошибка (авто)"
            
            self.backup_tree.insert("", "end", values=(date_str, file_path, status))

# --- Класс, описывающий создание фрейма О ПРОГРАММЕ ---
class About(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        testLbl = ttk.Label(self, background='#35A7FF')
        testLbl.pack(expand=True, fill='both')
        self.place(relx=0.3, y=0, relwidth=0.7, relheight=1)
        self.create_widgets()

    def create_widgets(self):
        # Основной контейнер
        container = ttk.Frame(self)
        container.place(relx=0.5, rely=0.5, anchor='center', width=400, height=200)
        
        # Заголовок
        ttk.Label(container, 
                 text="О программе", 
                 font=('Arial', 16, 'bold')).pack(pady=(20, 30))
        
        # Информация о разработчике
        info_frame = ttk.Frame(container)
        info_frame.pack()
        
        ttk.Label(info_frame, 
                 text="Разработчик: Воробьев А.В.\n\n"
                      "Версия: 1.0.0",
                 font=('Arial', 12),
                 justify='center').pack()
        

# Создание экземпляра класса
if __name__ == "__main__":
    App('АРМ Космос', (1200, 800))